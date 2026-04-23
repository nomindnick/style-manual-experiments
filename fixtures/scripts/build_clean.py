"""Build the canonical clean fixture (`fixtures/clean.docx`).

This is the Westlake USD v. Apex Construction demurrer, promoted from
`build_westlake.py` after Wave-B verification + multi-pass fix-up. It serves
as the Phase-0.5 reference for "what a fully LS-compliant draft looks like."
The Phase-1 rules should produce zero findings against this document.

One of three Phase-0.5 clean fixtures for the style-manual rule checker. This
script is the source of truth for the document — the .docx is produced by
running this file. See `fixtures/DRAFT_BRIEF.md` and
`fixtures/seed-citations.verified.md` for the spec and the allowed citations.

LS Style Manual conformance is enforced at the run level:
- Times New Roman 12pt on every run (caption, body, headings, footnotes).
- Exactly-24pt line spacing on body paragraphs.
- 0.5" first-line indent on body paragraphs (not headings, not block quotes).
- Two spaces between sentences and after colons (built into the body text).
- Oxford comma throughout.
- Case names italicized; LS short-cite form (no `supra` for cases).
- `§` / `subd.` / `¶` inside parens; spelled out in running text.
- `id.` (including the period) italicized; no `ibid.`.
- Block quote indented 0.5" both margins, single-spaced, no surrounding quotes.
- Footnote markers follow terminal punctuation.

Run from repo root: `.venv/bin/python fixtures/scripts/build_clean.py`.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(12)
BODY_LINE_SPACING = Pt(24)
BODY_FIRST_LINE_INDENT = Inches(0.5)
BLOCK_INDENT = Inches(0.5)

OUTPUT_PATH = Path("fixtures/clean.docx")


# ---------------------------------------------------------------------------
# Low-level run / paragraph helpers
# ---------------------------------------------------------------------------


def _style_run(run) -> None:
    """Force TNR 12pt on every run (including for east-asian fallback)."""
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    rpr = run._element.get_or_add_rPr()
    r_fonts = rpr.find(qn("w:rFonts"))
    if r_fonts is None:
        from docx.oxml import OxmlElement

        r_fonts = OxmlElement("w:rFonts")
        rpr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), FONT_NAME)
    r_fonts.set(qn("w:hAnsi"), FONT_NAME)
    r_fonts.set(qn("w:cs"), FONT_NAME)
    r_fonts.set(qn("w:eastAsia"), FONT_NAME)


def _apply_body_format(paragraph, first_line: bool = True) -> None:
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = BODY_LINE_SPACING
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if first_line:
        pf.first_line_indent = BODY_FIRST_LINE_INDENT


def add_run(paragraph, text: str, *, italic: bool = False, bold: bool = False):
    run = paragraph.add_run(text)
    run.italic = italic
    run.bold = bold
    _style_run(run)
    return run


# ---------------------------------------------------------------------------
# Rich-text body paragraph builder
#
# `segments` is a list of (text, style) tuples where style is one of
# "", "i", "b", "bi". This lets us interleave italic case names inside a
# single body paragraph without losing TNR 12pt or spacing.
# ---------------------------------------------------------------------------


def add_body(doc, segments, *, first_line: bool = True, align=None):
    p = doc.add_paragraph()
    _apply_body_format(p, first_line=first_line)
    if align is not None:
        p.alignment = align
    for text, style in segments:
        add_run(
            p,
            text,
            italic=("i" in style),
            bold=("b" in style),
        )
    return p


def add_plain(doc, text: str, *, first_line: bool = True, align=None):
    return add_body(doc, [(text, "")], first_line=first_line, align=align)


def add_heading(doc, text: str, *, center: bool = False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = BODY_LINE_SPACING
    pf.space_before = Pt(12)
    pf.space_after = Pt(0)
    pf.first_line_indent = Inches(0)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, text, bold=True)
    return p


def add_block_quote(doc, text: str, citation: str | None = None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.left_indent = BLOCK_INDENT
    pf.right_indent = BLOCK_INDENT
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.first_line_indent = Inches(0)
    add_run(p, text)
    if citation is not None:
        cp = doc.add_paragraph()
        cpf = cp.paragraph_format
        cpf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        cpf.line_spacing = BODY_LINE_SPACING
        cpf.space_before = Pt(0)
        cpf.space_after = Pt(0)
        cpf.first_line_indent = Inches(0)
        add_run(cp, citation)
    return p


def add_blank(doc):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = BODY_LINE_SPACING
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    add_run(p, "")
    return p


# ---------------------------------------------------------------------------
# Caption
# ---------------------------------------------------------------------------


def build_caption(doc) -> None:
    # Attorney-of-record block (top-left, single-spaced-ish)
    attorney_lines = [
        "Margaret E. Calderón, Esq. (SBN 248391)",
        "Julian T. Rhee, Esq. (SBN 312047)",
        "LOZANO SMITH",
        "2000 Crow Canyon Place, Suite 200",
        "San Ramon, CA 94583",
        "Telephone: (925) 302-2000",
        "Facsimile: (925) 302-2010",
        "Email: mcalderon@lozanosmith.com",
        "",
        "Attorneys for Defendant",
        "WESTLAKE UNIFIED SCHOOL DISTRICT",
    ]
    for line in attorney_lines:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = BODY_LINE_SPACING
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.first_line_indent = Inches(0)
        add_run(p, line)

    add_blank(doc)

    # Court name (centered, bold)
    court_lines = [
        "SUPERIOR COURT OF THE STATE OF CALIFORNIA",
        "FOR THE COUNTY OF ORANGE",
    ]
    for line in court_lines:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = BODY_LINE_SPACING
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.first_line_indent = Inches(0)
        add_run(p, line, bold=True)

    add_blank(doc)

    # Two-column-style caption done as a simple table for alignment
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    # Column widths: left ~3.5", right ~3.0"
    left_cell, right_cell = table.rows[0].cells
    left_cell.width = Inches(3.5)
    right_cell.width = Inches(3.0)

    left_lines = [
        ("APEX CONSTRUCTION, INC., a California", ""),
        ("corporation,", ""),
        ("", ""),
        ("                    Plaintiff,", ""),
        ("", ""),
        ("v.", ""),
        ("", ""),
        ("WESTLAKE UNIFIED SCHOOL DISTRICT;", ""),
        ("and DOES 1 through 25, inclusive,", ""),
        ("", ""),
        ("                    Defendants.", ""),
    ]
    right_lines = [
        ("Case No. 30-2025-01318742-CU-BC-CJC", "b"),
        ("", ""),
        ("DEMURRER TO PLAINTIFF'S COMPLAINT;", "b"),
        ("MEMORANDUM OF POINTS AND", "b"),
        ("AUTHORITIES IN SUPPORT THEREOF", "b"),
        ("", ""),
        ("Date:          June 18, 2026", ""),
        ("Time:          9:00 a.m.", ""),
        ("Dept.:         C-14", ""),
        ("Judge:         Hon. Teresa R. Halvorsen", ""),
        ("", ""),
        ("Action Filed:  February 3, 2026", ""),
        ("Trial Date:    None Set", ""),
    ]

    def _fill_cell(cell, lines):
        # Remove the auto-created empty paragraph
        cell._tc.remove(cell.paragraphs[0]._p)
        for text, style in lines:
            p = cell.add_paragraph()
            pf = p.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = BODY_LINE_SPACING
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.first_line_indent = Inches(0)
            add_run(p, text, bold=("b" in style), italic=("i" in style))

    _fill_cell(left_cell, left_lines)
    _fill_cell(right_cell, right_lines)

    add_blank(doc)


# ---------------------------------------------------------------------------
# Body — notice, memorandum, argument
# ---------------------------------------------------------------------------


def build_notice(doc) -> None:
    add_heading(
        doc,
        "NOTICE OF DEMURRER AND DEMURRER TO PLAINTIFF'S COMPLAINT",
        center=True,
    )
    add_heading(
        doc, "TO PLAINTIFF AND TO PLAINTIFF'S COUNSEL OF RECORD:", center=False
    )

    add_plain(
        doc,
        "PLEASE TAKE NOTICE that on June 18, 2026, at 9:00 a.m., or as soon "
        "thereafter as the matter may be heard, in Department C-14 of the "
        "above-entitled Court, located at 700 Civic Center Drive West, Santa "
        "Ana, California 92701, defendant Westlake Unified School District "
        "(\"District\") will, and hereby does, demur to the Complaint of "
        "plaintiff Apex Construction, Inc. (\"Plaintiff\") on the grounds "
        "set forth below.  This demurrer is based on this Notice, the "
        "accompanying Memorandum of Points and Authorities, the concurrently "
        "filed Declaration of Margaret E. Calderón regarding meet-and-confer "
        "efforts under Code of Civil Procedure section 430.41, the "
        "pleadings and records on file, and such further evidence and "
        "argument as may be presented at the hearing.",
    )

    add_plain(
        doc,
        "The District demurs to each cause of action in the Complaint on "
        "the grounds that none states facts sufficient to constitute a "
        "cause of action against the District, pursuant to Code of Civil "
        "Procedure section 430.10, subdivision (e), and is uncertain within "
        "the meaning of Code of Civil Procedure section 430.10, subdivision "
        "(f).  Specifically:",
    )

    grounds = [
        "1.  The first cause of action for breach of contract fails because "
        "the alleged \"extra work\" was never approved by the District's "
        "Board of Trustees (\"Board\") by formal action, as Education Code "
        "sections 17604 and 17605 and Public Contract Code section "
        "20118.4 require; absent board approval and a written change "
        "order, no enforceable contract amendment exists.",
        "2.  The second cause of action for quantum meruit fails because a "
        "quasi-contractual recovery is not available against a public "
        "entity for work performed outside the scope of a board-approved "
        "public contract.",
        "3.  The third cause of action for promissory estoppel fails "
        "because equitable estoppel may not be asserted against a public "
        "entity where it would nullify a statutory policy — here, the "
        "policy requiring board approval of school-district contracts.",
    ]
    for g in grounds:
        add_plain(doc, g)

    add_plain(
        doc,
        "This demurrer follows the parties' meet-and-confer efforts as "
        "required by Code of Civil Procedure section 430.41, subdivision "
        "(a).  The parties were unable to resolve the pleading defects "
        "identified herein.",
    )


def build_memorandum(doc) -> None:
    add_heading(doc, "MEMORANDUM OF POINTS AND AUTHORITIES", center=True)

    # I. INTRODUCTION
    add_heading(doc, "I.  INTRODUCTION")
    add_body(
        doc,
        [
            (
                "Plaintiff is a contractor that performed public works for "
                "the District under a competitively bid construction "
                "contract.  Plaintiff now seeks to recover more than "
                "$1.4 million in claimed \"extras\" that were never approved "
                "by the District's Board and were never memorialized in a "
                "written change order.  Plaintiff's theory is, in essence, "
                "that because a construction manager and the District's "
                "on-site project inspector allegedly told Plaintiff to "
                "proceed with additional electrical and HVAC work, the "
                "District is now bound to pay for it.  The law is to the "
                "contrary.",
                "",
            ),
        ],
    )
    add_body(
        doc,
        [
            (
                "California has long imposed strict, mandatory procedures on "
                "how public school districts enter into and modify contracts.  "
                "Those procedures exist to protect taxpayers and to prevent "
                "precisely the kind of after-the-fact claim Plaintiff now "
                "presses.  Because each of Plaintiff's three causes of "
                "action — breach of contract, quantum meruit, and "
                "promissory estoppel — depends on enforcing obligations "
                "that the District's Board never approved, none states a "
                "cognizable claim against a California public school "
                "district.  The demurrer should be sustained without leave "
                "to amend.",
                "",
            ),
        ],
    )

    # II. STATEMENT OF FACTS
    add_heading(doc, "II.  STATEMENT OF RELEVANT ALLEGATIONS")
    add_body(
        doc,
        [
            (
                "For purposes of this demurrer only, the District accepts "
                "the well-pleaded factual allegations of the Complaint as "
                "true.  (",
                "",
            ),
            ("Blank v. Kirwan", "i"),
            (
                " (1985) 39 Cal.3d 311, 318.)  The Complaint alleges as "
                "follows.",
                "",
            ),
        ],
    )
    add_body(
        doc,
        [
            (
                "In September 2024, the District issued an invitation for "
                "bids for the Westlake High School Science Classroom "
                "Modernization Project, a phased renovation of five "
                "existing science classrooms funded through a Proposition 39 "
                "bond measure.  (Compl., ¶¶ 6–8.)  The project was bid under "
                "the competitive-bidding requirements of Public Contract "
                "Code section 20111.  (Compl., ¶ 7.)  Plaintiff submitted "
                "the lowest responsive bid and, on November 14, 2024, the "
                "Board awarded Plaintiff a fixed-price construction contract "
                "in the amount of $4,217,450 (the \"Contract\").  (Compl., "
                "¶¶ 9–10.)  The Board's award is reflected in Resolution "
                "No. 2024-27.  (Compl., ¶ 10.)",
                "",
            ),
        ],
    )
    add_body(
        doc,
        [
            (
                "The Contract incorporates the District's standard General "
                "Conditions, including a written change-order clause.  "
                "(Compl., ¶ 12, Ex. A [General Conditions, § 9.1].)  That "
                "clause, consistent with Public Contract Code section "
                "20118.4, provides that the contract sum may be modified "
                "only by a written change order executed by an authorized "
                "District representative and, where required, approved by "
                "the Board.  (Compl., Ex. A [General Conditions, § 9.1].)",
                "",
            ),
        ],
    )
    add_body(
        doc,
        [
            (
                "Beginning in late March 2025, Plaintiff alleges, it "
                "encountered field conditions in the existing classrooms "
                "that required additional electrical rework and HVAC "
                "modifications.  (Compl., ¶ 15.)  Plaintiff alleges that "
                "the District's construction manager, Meridian CM "
                "Services, Inc., and the District's on-site project "
                "inspector, Paul Brennan, verbally directed Plaintiff to "
                "proceed with the additional work.  (Compl., ¶¶ 16–17.)  "
                "Plaintiff further alleges that, in a series of emails "
                "between April 4, 2025 and July 22, 2025, Mr. Brennan "
                "confirmed that \"the extras would be paid.\"  (Compl., "
                "¶¶ 18–21, Exs. B–E.)  The Complaint alleges, in pertinent "
                "part:",
                "",
            ),
        ],
    )
    add_block_quote(
        doc,
        "Throughout the performance of the extra work, Plaintiff relied "
        "on the repeated written confirmations of the District's project "
        "inspector, Paul Brennan, that the extras would be paid in the "
        "ordinary course, that no further written approval was required "
        "because the field changes were within the inspector's scope of "
        "authority, and that Plaintiff should continue performance "
        "without delay so as not to jeopardize the District's occupancy "
        "schedule for the 2025–2026 academic year.",
        citation="(Compl., ¶ 20.)",
    )
    add_body(
        doc,
        [
            (
                "Plaintiff concedes, as it must, that the Board never "
                "approved the alleged extras by formal action and that no "
                "written change order was ever executed for any of the "
                "additional work.  (Compl., ¶¶ 24–25.)  Plaintiff "
                "nevertheless submitted a final invoice on November 3, "
                "2025, in the amount of $1,412,880, which the District "
                "declined to pay.  (Compl., ¶¶ 27–28.)  This action "
                "followed.",
                "",
            ),
        ],
    )
    add_body(
        doc,
        [
            (
                "The Complaint pleads three causes of action:  (1) breach "
                "of the Contract based on non-payment of the alleged "
                "extras; (2) quantum meruit for the reasonable value of "
                "the same field-directed work; and (3) promissory estoppel "
                "based on Mr. Brennan's emails.  (Compl., ¶¶ 30–58.)  Each "
                "fails on the face of the Complaint.",
                "",
            ),
        ],
    )

    # III. LEGAL STANDARD
    add_heading(doc, "III.  LEGAL STANDARD")
    add_body(
        doc,
        [
            (
                "A demurrer tests the legal sufficiency of a complaint.  "
                "(",
                "",
            ),
            ("Blank v. Kirwan", "i"),
            (
                " (1985) 39 Cal.3d 311, 318.)  In reviewing a demurrer, "
                "the court accepts as true all well-pleaded material facts, "
                "but not contentions, deductions, or conclusions of law.  "
                "(",
                "",
            ),
            ("Zelig v. County of Los Angeles", "i"),
            (
                " (2002) 27 Cal.4th 1112, 1126.)  The demurrer reaches "
                "only the face of the complaint and those matters that "
                "are subject to judicial notice.  (",
                "",
            ),
            ("Aubry v. Tri-City Hospital Dist.", "i"),
            (
                " (1992) 2 Cal.4th 962, 967.)  If the complaint fails to "
                "state facts sufficient to constitute a cause of action, "
                "the court must determine whether there is a reasonable "
                "possibility the defect can be cured by amendment; the "
                "plaintiff bears the burden of demonstrating such a "
                "possibility.  (",
                "",
            ),
            ("Schifando v. City of Los Angeles", "i"),
            (" (2003) 31 Cal.4th 1074, 1081.)", ""),
        ],
    )
    add_body(
        doc,
        [
            (
                "Where, as here, the defect is a legal bar to recovery "
                "against a public entity — not a curable pleading defect — "
                "leave to amend should be denied.  (",
                "",
            ),
            ("Blank", "i"),
            (
                ", 39 Cal.3d at p. 318; ",
                "",
            ),
            ("Schifando", "i"),
            (", 31 Cal.4th at p. 1081.)", ""),
        ],
    )

    # IV. ARGUMENT
    add_heading(doc, "IV.  ARGUMENT")

    # A. Breach of contract
    add_heading(
        doc,
        "A.  The First Cause of Action for Breach of Contract Fails Because "
        "the Alleged \"Extras\" Were Never Approved by the Board or "
        "Memorialized in a Written Change Order.",
    )
    add_body(
        doc,
        [
            (
                "California imposes mandatory procedures on how school "
                "districts enter into and amend contracts.  Education Code "
                "section 17604 provides that no contract binds a school "
                "district unless it is approved or ratified by the "
                "governing board, and Education Code section 17605 "
                "authorizes the board to delegate certain contracting "
                "authority only under defined conditions.  These "
                "provisions are strict and jurisdictional:  compliance is a "
                "prerequisite to the existence of any enforceable "
                "obligation on the public entity's part.  (",
                "",
            ),
            ("G.L. Mezzetta, Inc. v. City of American Canyon", "i"),
            (
                " (2000) 78 Cal.App.4th 1087, 1094–1095; ",
                "",
            ),
            ("First Street Plaza Partners v. City of Los Angeles", "i"),
            (" (1998) 65 Cal.App.4th 650, 669–670.)", ""),
        ],
    )
    add_body(
        doc,
        [
            (
                "For school-district public works in particular, Public "
                "Contract Code section 20118.4 requires that changes to "
                "the contract be made by written change order.  Read "
                "together with Education Code sections 17604 and 17605, "
                "the statutes establish an unbroken rule:  absent a board-"
                "approved, written change order, there is no contract "
                "amendment.  Any purported oral or informal modification — "
                "however well-intentioned the field personnel who made "
                "it — is a legal nullity.  (",
                "",
            ),
            ("P&D Consultants, Inc. v. City of Carlsbad", "i"),
            (
                " (2010) 190 Cal.App.4th 1332, 1339–1341 [written change-"
                "order requirements in public contracts strictly enforced]; ",
                "",
            ),
            ("Amelco Electric v. City of Thousand Oaks", "i"),
            (
                " (2002) 27 Cal.4th 228, 234–235.)",
                "",
            ),
        ],
    )
    add_body(
        doc,
        [
            (
                "The rationale is familiar.  The essentials of a contract "
                "under California law include the consent of the parties "
                "and a lawful object.  (Civ. Code, § 1550.)  Consent must "
                "be mutual and communicated by each party to the other.  "
                "(Civ. Code, § 1565.)  When the Legislature has specified "
                "the manner in which a public entity may give its "
                "consent — here, by board action and a written change "
                "order — a contractor who proceeds on the say-so of an "
                "inspector or construction manager does so at its own "
                "risk.  (",
                "",
            ),
            ("Reams v. Cooley", "i"),
            (
                " (1915) 171 Cal. 150, 153–154.)  The public entity's "
                "capacity to contract is confined to the channels the "
                "Legislature has authorized; anything outside those "
                "channels is not merely voidable but void ",
                "",
            ),
            ("ab initio", "i"),
            (".  (", ""),
            ("First Street Plaza Partners", "i"),
            (", 65 Cal.App.4th at pp. 669–670.)", ""),
        ],
    )
    add_body(
        doc,
        [
            (
                "Plaintiff's Complaint concedes the dispositive facts.  "
                "Plaintiff acknowledges that (i) the Board awarded the "
                "Contract by formal action in Resolution No. 2024-27; "
                "(ii) the Board never approved any increase to the "
                "contract sum for the alleged extras; and (iii) no "
                "written change order was ever executed for any of the "
                "additional work.  (Compl., ¶¶ 10, 24–25.)  Those "
                "concessions are fatal.  The alleged oral directions and "
                "emails from the District's on-site inspector and the "
                "construction manager cannot, as a matter of law, "
                "substitute for the statutorily required board approval "
                "and written change order.  (Ed. Code, §§ 17604, 17605; "
                "Pub. Cont. Code, § 20118.4, subd. (a); ",
                "",
            ),
            ("P&D Consultants", "i"),
            (", 190 Cal.App.4th at p. 1341.)", ""),
        ],
    )
    add_body(
        doc,
        [
            (
                "Nor may Plaintiff rescue the breach-of-contract claim by "
                "invoking the implied covenant of good faith and fair "
                "dealing.  The implied covenant cannot be used to impose "
                "on a party substantive duties or obligations beyond "
                "those expressly provided for in the contract.  (",
                "",
            ),
            ("Carma Developers (Cal.), Inc. v. Marathon Development "
             "California, Inc.", "i"),
            (
                " (1992) 2 Cal.4th 342, 374; ",
                "",
            ),
            ("Guz v. Bechtel National, Inc.", "i"),
            (
                " (2000) 24 Cal.4th 317, 349–350.)  Because the express "
                "Contract required a written, board-approved change order "
                "for any modification of the contract sum, an implied "
                "covenant cannot be enlisted to dispense with that "
                "requirement.  (",
                "",
            ),
            ("Guz", "i"),
            (", 24 Cal.4th at pp. 349–350.)", ""),
        ],
    )
    add_body(
        doc,
        [
            (
                "The first cause of action should therefore be sustained "
                "without leave to amend.  Plaintiff cannot, consistent "
                "with its own factual concessions, allege facts showing "
                "that the Board approved the extras or that a written "
                "change order ever issued.  (",
                "",
            ),
            ("Blank", "i"),
            (
                ", 39 Cal.3d at p. 318.)",
                "",
            ),
        ],
    )

    # B. Quantum meruit
    add_heading(
        doc,
        "B.  The Second Cause of Action for Quantum Meruit Fails Because "
        "Quasi-Contractual Recovery Is Unavailable Against a Public Entity "
        "for Work Performed Outside a Board-Approved Contract.",
    )
    add_body(
        doc,
        [
            (
                "Plaintiff's second cause of action attempts to recover "
                "the \"reasonable value\" of the same field-directed "
                "extras through a theory of quantum meruit.  California "
                "law has foreclosed that path for more than a century.  "
                "Where a statute prescribes the manner in which a public "
                "entity's contract must be formed, a contractor cannot "
                "circumvent the statute by repackaging its claim in "
                "quasi-contract.  (",
                "",
            ),
            ("Miller v. McKinnon", "i"),
            (
                " (1942) 20 Cal.2d 83, 88–89; ",
                "",
            ),
            ("Reams", "i"),
            (
                ", 171 Cal. at pp. 153–154.)",
                "",
            ),
        ],
    )
    add_body(
        doc,
        [
            (
                "The rule has been applied consistently to bar quantum "
                "meruit claims premised on work that would have required "
                "board approval under the governing public-contracting "
                "statute.  In ",
                "",
            ),
            ("Katsura v. City of San Buenaventura", "i"),
            (
                " (2007) 155 Cal.App.4th 104, the Court of Appeal held "
                "that a plaintiff who had performed services without "
                "complying with the city's contracting requirements "
                "could not recover in quantum meruit because permitting "
                "such a recovery would circumvent the very statutes the "
                "Legislature enacted to protect the public fisc.  (",
                "",
            ),
            ("Id.", "i"),
            (
                " at pp. 108–110.)  The ",
                "",
            ),
            ("Katsura", "i"),
            (
                " court relied on the foundational rule of ",
                "",
            ),
            ("Miller v. McKinnon", "i"),
            (
                ":  where a contract with a public entity is void for want "
                "of statutory compliance, quasi-contractual recovery is "
                "also unavailable, because permitting a quasi-contractual "
                "recovery would give contractors an easy route to "
                "circumvent the statutory requirements the Legislature "
                "imposed on public-entity contracting.¹  (",
                "",
            ),
            ("Katsura", "i"),
            (
                ", 155 Cal.App.4th at p. 109; ",
                "",
            ),
            ("Miller", "i"),
            (", 20 Cal.2d at p. 89.)", ""),
        ],
    )
    # Footnote 1 — added below in `build_footnote_1`
    add_body(
        doc,
        [
            (
                "More recently, in ",
                "",
            ),
            ("G.L. Mezzetta, Inc. v. City of American Canyon", "i"),
            (
                ", the Court of Appeal confirmed that strict compliance "
                "with public-contracting statutes is jurisdictional and "
                "that a contractor's equitable or quasi-contractual "
                "claims rise and fall with the validity of the underlying "
                "contract.  (",
                "",
            ),
            ("G.L. Mezzetta", "i"),
            (
                ", 78 Cal.App.4th at pp. 1094–1095.)  The rule is not a "
                "technicality; it is the Legislature's considered "
                "allocation of risk.  Contractors deal with public "
                "entities on notice of the statutes that govern the "
                "agency's capacity, and they are charged with knowledge "
                "of those statutes.  (",
                "",
            ),
            ("First Street Plaza Partners", "i"),
            (", 65 Cal.App.4th at p. 670.)", ""),
        ],
    )
    add_body(
        doc,
        [
            (
                "The same analysis applies here.  Plaintiff's quantum "
                "meruit claim seeks compensation for precisely the work "
                "that could have been lawfully compensated only through a "
                "board-approved change order under Education Code "
                "sections 17604 and 17605 and Public Contract Code "
                "section 20118.4.  Because the Board never approved the "
                "extras and no change order ever issued, there is no "
                "valid public contract — and, under ",
                "",
            ),
            ("Miller", "i"),
            (
                ", ",
                "",
            ),
            ("Katsura", "i"),
            (
                ", and ",
                "",
            ),
            ("G.L. Mezzetta", "i"),
            (
                ", no quasi-contractual recovery either.  Plaintiff's "
                "attempt to invoke the abandonment doctrine or a \"total "
                "cost\" measure of damages is likewise unavailable:  the "
                "Supreme Court squarely rejected those theories as "
                "applied to public-entity contractors in ",
                "",
            ),
            ("Amelco Electric v. City of Thousand Oaks", "i"),
            (
                ", 27 Cal.4th at pp. 234–237.",
                "",
            ),
        ],
    )
    add_body(
        doc,
        [
            (
                "The second cause of action should therefore be sustained "
                "without leave to amend.  No amendment can cure the "
                "absence of a board-approved contract for the extras; "
                "that absence is admitted on the face of the Complaint.  "
                "(Compl., ¶¶ 24–25; see ",
                "",
            ),
            ("Schifando", "i"),
            (", 31 Cal.4th at p. 1081.)", ""),
        ],
    )

    # C. Promissory estoppel
    add_heading(
        doc,
        "C.  The Third Cause of Action for Promissory Estoppel Fails "
        "Because Equitable Estoppel Cannot Be Invoked to Nullify the "
        "Statutory Policy Requiring Board Approval of School-District "
        "Contracts.",
    )
    add_body(
        doc,
        [
            (
                "Plaintiff's third cause of action asks the Court to "
                "treat Mr. Brennan's emails — the same communications "
                "that could not amend the Contract under Education Code "
                "sections 17604 and 17605 or Public Contract Code "
                "section 20118.4 — as binding promises enforceable "
                "through promissory estoppel.  The law does not permit "
                "that result.",
                "",
            ),
        ],
    )
    add_body(
        doc,
        [
            (
                "The California Supreme Court has long held that "
                "equitable estoppel may be invoked against the government "
                "only in genuinely exceptional situations, and never "
                "where its application would nullify a statutory policy.  (",
                "",
            ),
            ("City of Long Beach v. Mansell", "i"),
            (
                " (1970) 3 Cal.3d 462, 493.)  That limitation is "
                "categorical.  In ",
                "",
            ),
            ("Lentz v. McMahon", "i"),
            (
                " (1989) 49 Cal.3d 393, the Supreme Court reaffirmed that "
                "estoppel is unavailable against a public entity where it "
                "would frustrate a strong public policy adopted for the "
                "benefit of the public.  (",
                "",
            ),
            ("Id.", "i"),
            (
                " at pp. 399–400.)  The Court of Appeal followed the same "
                "rule in ",
                "",
            ),
            ("Janis v. California State Lottery Com.", "i"),
            (
                " (1998) 68 Cal.App.4th 824, 830–831, observing that "
                "estoppel claims against public entities are narrowly "
                "confined and may not be used to override statutory "
                "mandates.",
                "",
            ),
        ],
    )
    add_body(
        doc,
        [
            (
                "The policy at stake here is exactly the kind the "
                "Supreme Court had in mind.  The Legislature has "
                "determined that school-district contracts may be "
                "formed and modified only through formal board action "
                "and, for public works, a written change order.  (Ed. "
                "Code, §§ 17604, 17605; Pub. Cont. Code, § 20118.4.)  "
                "That requirement is not a mere formality; it is the "
                "central mechanism by which the Legislature protects "
                "public funds and ensures democratic accountability for "
                "school-district spending.  A rule permitting a "
                "contractor to bind a school district through the "
                "informal assurances of an on-site inspector — bypassing "
                "the board entirely — would hollow out the statutory "
                "scheme.  That is precisely the result ",
                "",
            ),
            ("Mansell", "i"),
            (
                ", ",
                "",
            ),
            ("Lentz", "i"),
            (
                ", and ",
                "",
            ),
            ("Janis", "i"),
            (" forbid.", ""),
        ],
    )
    add_body(
        doc,
        [
            (
                "The rule applies with special force where, as here, the "
                "public official whose conduct is said to give rise to "
                "the estoppel lacked statutory authority to bind the "
                "entity in the first place.  (",
                "",
            ),
            ("Air Quality Products, Inc. v. State of California", "i"),
            (
                " (1979) 96 Cal.App.3d 340, 348–349 [no estoppel where "
                "agent exceeded statutory authority].)  The District's "
                "on-site project inspector is not the Board; he is a "
                "field employee whose authority is delimited by the "
                "Contract and by the inspection statutes governing "
                "school construction.  Nothing in the Complaint — and "
                "nothing the Complaint could allege consistent with its "
                "existing averments — places Mr. Brennan within the "
                "narrow delegation authorized by Education Code section "
                "17605.  His assurances, however sincere, could not and "
                "cannot bind the District.  (",
                "",
            ),
            ("Air Quality Products", "i"),
            (", 96 Cal.App.3d at pp. 348–349.)", ""),
        ],
    )
    add_body(
        doc,
        [
            (
                "Finally, Plaintiff cannot plead around these authorities "
                "by labeling its claim \"promissory\" rather than "
                "\"equitable\" estoppel.  The Supreme Court and Courts "
                "of Appeal have applied the same limitation to both "
                "doctrines when asserted against public entities, "
                "because the underlying concern — that estoppel not be "
                "used to dispense with statutory prerequisites to public "
                "liability — is the same.  (",
                "",
            ),
            ("Lentz", "i"),
            (
                ", 49 Cal.3d at pp. 399–400; ",
                "",
            ),
            ("Janis", "i"),
            (
                ", 68 Cal.App.4th at pp. 830–831; see also Gov. Code, "
                "§ 815 [public-entity liability exists only as provided "
                "by statute].)",
                "",
            ),
        ],
    )
    add_body(
        doc,
        [
            (
                "The third cause of action should therefore be sustained "
                "without leave to amend.",
                "",
            ),
        ],
    )

    # V. CONCLUSION
    add_heading(doc, "V.  CONCLUSION")
    add_body(
        doc,
        [
            (
                "Each of Plaintiff's three causes of action depends on "
                "enforcing obligations that the District's Board never "
                "approved and that no written change order ever "
                "memorialized.  California law does not permit a "
                "contractor to impose such obligations on a public "
                "school district through oral directions, informal "
                "emails, or a repackaged theory of equity.  The "
                "District respectfully requests that the Court sustain "
                "its demurrer to the first, second, and third causes of "
                "action of the Complaint, without leave to amend, and "
                "grant such further relief as the Court deems just and "
                "proper.",
                "",
            ),
        ],
    )


def build_signature(doc) -> None:
    add_blank(doc)
    add_plain(doc, "Dated: April 23, 2026", first_line=False)
    add_plain(doc, "", first_line=False)
    add_plain(doc, "LOZANO SMITH", first_line=False)
    add_plain(doc, "", first_line=False)
    add_plain(doc, "By:  /s/ Margaret E. Calderón", first_line=False)
    add_plain(doc, "     Margaret E. Calderón", first_line=False)
    add_plain(doc, "     Julian T. Rhee", first_line=False)
    add_plain(doc, "     Attorneys for Defendant", first_line=False)
    add_plain(
        doc,
        "     WESTLAKE UNIFIED SCHOOL DISTRICT",
        first_line=False,
    )


def build_footnote_section(doc) -> None:
    """Render footnotes as simple end-of-document single-spaced paragraphs.

    python-docx does not have native footnote support; LS-FN-02/03 specify
    TNR 12pt single-spaced text, which we reproduce here. The marker in
    the body text is the ¹ character; this paragraph begins with the same
    marker.
    """
    # Separator
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.first_line_indent = Inches(0)
    pf.space_before = Pt(12)
    pf.space_after = Pt(0)
    add_run(p, "_______________________")

    fn = doc.add_paragraph()
    fpf = fn.paragraph_format
    fpf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    fpf.first_line_indent = Inches(0)
    fpf.space_before = Pt(0)
    fpf.space_after = Pt(0)
    add_run(
        fn,
        "¹ The ",
    )
    add_run(fn, "Miller v. McKinnon", italic=True)
    add_run(
        fn,
        " rule was already well-established when the Legislature "
        "enacted the modern Public Contract Code; the statutory "
        "scheme has been construed against the backdrop of that rule.  "
        "(See ",
    )
    add_run(fn, "Reams v. Cooley", italic=True)
    add_run(fn, ", 171 Cal. at pp. 153–154.)")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------


def main() -> None:
    doc = Document()

    # Set default style to TNR 12pt as well (belt and suspenders).
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = FONT_SIZE
    rpr = normal.element.get_or_add_rPr()
    r_fonts = rpr.find(qn("w:rFonts"))
    if r_fonts is None:
        from docx.oxml import OxmlElement

        r_fonts = OxmlElement("w:rFonts")
        rpr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), FONT_NAME)
    r_fonts.set(qn("w:hAnsi"), FONT_NAME)
    r_fonts.set(qn("w:cs"), FONT_NAME)
    r_fonts.set(qn("w:eastAsia"), FONT_NAME)

    # 1" margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    build_caption(doc)
    build_notice(doc)
    build_memorandum(doc)
    build_signature(doc)
    build_footnote_section(doc)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"Wrote {OUTPUT_PATH} ({OUTPUT_PATH.stat().st_size} bytes).")


if __name__ == "__main__":
    main()
