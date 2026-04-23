"""Build `fixtures/realistic-mixed.docx` — the REALISTIC MIXED corruption fixture.

Wave-C corruption pass over the clean Wave-A Sierra Vista demurrer. The vibe is
"a real LS draft that almost made it through proofreading." Five intentional
violations, all plausible typos planted deep in the arguments (mid-paragraph,
not at paragraph starts). The intentional violations are recorded in the
sidecar ground-truth file `fixtures/realistic-mixed.violations.json`.

Source: Sierra Vista clean Wave-A draft (now superseded; audit trail in
`fixtures/realistic-mixed.SOURCE_NOTES.md` and `.SOURCE_VIOLATIONS.md`).
Corruption brief: `fixtures/CORRUPT_BRIEF.md`.

Intentional violations injected (see JSON for exact snippets + descriptions):

  1. LS-SP-02 (error)    ¶62 — single space between sentences (was two) at
     "...154.) Contractors..." in the account-stated argument.
  2. LS-SP-02 (error)    ¶72 — single space between sentences (was two) at
     "...17605. Pre-bid representations..." in the declaratory-relief argument.
  3. LS-CITE-02 (error)  ¶60 — inside a parenthetical cite, `§ 20118.4` was
     rewritten to `section 20118.4` (section symbol replaced by word inside
     parens).
  4. LS-CAP-02 (warning) ¶81 — "bind the District despite..." was lowercased
     to "bind the district despite..." deep in the estoppel argument.
  5. LS-CITE-HAL (error) ¶82 — Air Quality Products page number transposed:
     `96 Cal.App.3d 340` → `96 Cal.App.3d 304`. The transposed cite does not
     resolve to the real case on CourtListener.

Every other citation in the document remains on
`fixtures/seed-citations.verified.md`.

Run from repo root:
    .venv/bin/python fixtures/scripts/build_realistic-mixed.py

This script is the source; the .docx is an artifact. Edits should be made
here, then the script re-run. Each injected violation is marked with a
`# CORRUPT(<rule_id>)` comment so a reader can find the planted-error
edits at a glance.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt


# ---------------------------------------------------------------------------
# Constants — fact-pattern specifics
# ---------------------------------------------------------------------------

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(12)
BODY_LINE_SPACING = Pt(24)
BODY_FIRST_LINE_INDENT = Inches(0.5)

COUNTY = "Tulare"
CASE_NO = "VCU 295418"
DEPT = "Department 7"
JUDGE = "Hon. Marisol D. Hernandez"
ACTION_FILED = "October 14, 2025"
TRIAL_DATE = "Not yet set"
HEARING_DATE = "June 18, 2026"
HEARING_TIME = "8:30 a.m."

DISTRICT_FULL = "Sierra Vista Unified School District"
DISTRICT_SHORT = "District"
PLAINTIFF_FULL = "Granite Peak Builders, Inc."
PLAINTIFF_SHORT = "Granite Peak"

ATTY_NAME = "Daniel R. Marchetti"
ATTY_SBN = "248917"
ATTY_EMAIL = "dmarchetti@lozanosmith.com"
ATTY_2_NAME = "Priya S. Ramanathan"
ATTY_2_SBN = "312045"
ATTY_2_EMAIL = "pramanathan@lozanosmith.com"
FIRM_NAME = "LOZANO SMITH"
FIRM_ADDRESS_1 = "One Capitol Mall, Suite 640"
FIRM_ADDRESS_2 = "Sacramento, CA 95814"
FIRM_PHONE = "Telephone: (916) 443-1800"
FIRM_FAX = "Facsimile: (916) 443-1801"

SIGNATURE_DATE = "April 23, 2026"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _style_run(run, *, italic: bool = False, bold: bool = False, underline: bool = False) -> None:
    """Apply TNR 12 pt plus optional italic/bold/underline to a run."""
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    # Set the east-Asian font too so Word doesn't substitute another face.
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), FONT_NAME)
    if italic:
        run.italic = True
    if bold:
        run.bold = True
    if underline:
        run.underline = True


def _apply_body_format(paragraph) -> None:
    """Exactly-24 pt line spacing, 0.5" first-line indent, justified."""
    pf = paragraph.paragraph_format
    pf.line_spacing = BODY_LINE_SPACING
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.first_line_indent = BODY_FIRST_LINE_INDENT
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _apply_plain_format(paragraph, *, alignment=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    """Exactly-24 pt line spacing, no first-line indent, alignment configurable."""
    pf = paragraph.paragraph_format
    pf.line_spacing = BODY_LINE_SPACING
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    paragraph.alignment = alignment


def add_body_paragraph(doc, segments):
    """Add a body paragraph composed of (text, formatting) segments.

    Each segment is either a plain string or a dict of the form
    ``{"text": "...", "italic": True}``. The paragraph gets standard body
    formatting: TNR 12, exactly-24 line spacing, 0.5" first-line indent,
    justified.
    """
    p = doc.add_paragraph()
    _apply_body_format(p)
    for seg in segments:
        if isinstance(seg, str):
            run = p.add_run(seg)
            _style_run(run)
        else:
            run = p.add_run(seg["text"])
            _style_run(
                run,
                italic=seg.get("italic", False),
                bold=seg.get("bold", False),
                underline=seg.get("underline", False),
            )
    return p


def add_heading(doc, text, *, level=1, center=True):
    """ALL-CAPS bold heading. Level 1 = Roman-numeral top heading."""
    p = doc.add_paragraph()
    _apply_plain_format(
        p,
        alignment=WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT,
    )
    run = p.add_run(text)
    _style_run(run, bold=True)
    return p


def add_blank(doc):
    p = doc.add_paragraph()
    _apply_plain_format(p)
    return p


def add_block_quote(doc, text, citation_segments):
    """A ≥50-word block quote: 0.5" both margins, single-spaced, no quotation
    marks. Citation goes on its own line below, at the left margin with a
    normal first-line indent (the citation sentence is in parens, period
    inside)."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.5)
    pf.right_indent = Inches(0.5)
    pf.first_line_indent = Inches(0)
    pf.line_spacing = 1.0
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    _style_run(run)

    cite_p = doc.add_paragraph()
    _apply_body_format(cite_p)
    cite_p.paragraph_format.first_line_indent = Inches(0)
    for seg in citation_segments:
        if isinstance(seg, str):
            run = cite_p.add_run(seg)
            _style_run(run)
        else:
            run = cite_p.add_run(seg["text"])
            _style_run(run, italic=seg.get("italic", False))
    return p


# ---------------------------------------------------------------------------
# Caption block
# ---------------------------------------------------------------------------


def build_caption(doc):
    # Attorneys of record, flush-left, single-spaced.
    atty_lines = [
        f"{ATTY_NAME}, Esq. (SBN {ATTY_SBN})",
        f"{ATTY_EMAIL}",
        f"{ATTY_2_NAME}, Esq. (SBN {ATTY_2_SBN})",
        f"{ATTY_2_EMAIL}",
        FIRM_NAME,
        FIRM_ADDRESS_1,
        FIRM_ADDRESS_2,
        FIRM_PHONE,
        FIRM_FAX,
        "",
        f"Attorneys for Defendant",
        f"{DISTRICT_FULL.upper()}",
    ]
    for line in atty_lines:
        p = doc.add_paragraph()
        _apply_plain_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        if line:
            run = p.add_run(line)
            _style_run(run)

    add_blank(doc)

    court_lines = [
        "SUPERIOR COURT OF THE STATE OF CALIFORNIA",
        f"COUNTY OF {COUNTY.upper()}",
    ]
    for line in court_lines:
        p = doc.add_paragraph()
        _apply_plain_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        run = p.add_run(line)
        _style_run(run, bold=True)

    add_blank(doc)

    # Caption grid — two columns. python-docx tables work fine here.
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    # Left col: parties / title. Right col: case info.
    table.columns[0].width = Inches(3.5)
    table.columns[1].width = Inches(3.5)

    left_cell = table.cell(0, 0)
    right_cell = table.cell(0, 1)

    # Left column — parties and title
    left_lines = [
        (PLAINTIFF_FULL + ",", False),
        ("", False),
        ("Plaintiff,", False),
        ("", False),
        ("v.", False),
        ("", False),
        (f"{DISTRICT_FULL}; and DOES 1 through 25, inclusive,", False),
        ("", False),
        ("Defendants.", False),
        ("", False),
        ("_______________________________________", False),
    ]
    first = True
    for text, bold in left_lines:
        if first:
            p = left_cell.paragraphs[0]
            first = False
        else:
            p = left_cell.add_paragraph()
        _apply_plain_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        if text:
            run = p.add_run(text)
            _style_run(run, bold=bold)

    # Right column — case no., title, hearing info
    right_content = [
        ("Case No. " + CASE_NO, True, False),
        ("", False, False),
        ("DEMURRER TO PLAINTIFF'S COMPLAINT;", True, False),
        ("MEMORANDUM OF POINTS AND AUTHORITIES", True, False),
        ("IN SUPPORT THEREOF", True, False),
        ("", False, False),
        (f"Date:  {HEARING_DATE}", False, False),
        (f"Time:  {HEARING_TIME}", False, False),
        (f"Dept.:  {DEPT}", False, False),
        (f"Judge:  {JUDGE}", False, False),
        ("", False, False),
        (f"Action Filed:  {ACTION_FILED}", False, False),
        (f"Trial Date:  {TRIAL_DATE}", False, False),
    ]
    first = True
    for text, bold, _ in right_content:
        if first:
            p = right_cell.paragraphs[0]
            first = False
        else:
            p = right_cell.add_paragraph()
        _apply_plain_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        if text:
            run = p.add_run(text)
            _style_run(run, bold=bold)

    add_blank(doc)


# ---------------------------------------------------------------------------
# Notice + demurrer
# ---------------------------------------------------------------------------


def build_notice_and_demurrer(doc):
    add_heading(doc, "NOTICE OF DEMURRER AND DEMURRER")
    add_blank(doc)

    add_body_paragraph(
        doc,
        [
            "TO PLAINTIFF AND ITS ATTORNEYS OF RECORD:",
        ],
    )
    add_body_paragraph(
        doc,
        [
            f"PLEASE TAKE NOTICE that on {HEARING_DATE}, at {HEARING_TIME}, "
            f"or as soon thereafter as the matter may be heard in {DEPT} of the above-entitled "
            f"Court, located at 221 South Mooney Boulevard, Visalia, California 93291, "
            f"Defendant ",
            {"text": DISTRICT_FULL, "text_plain": True},
            f" (“{DISTRICT_SHORT}”) will, and hereby does, demur to the Complaint of "
            f"Plaintiff ",
            {"text": PLAINTIFF_FULL, "text_plain": True},
            f" (“{PLAINTIFF_SHORT}” or “Plaintiff”) on the grounds set forth "
            f"below.",
        ],
    )
    add_body_paragraph(
        doc,
        [
            "This Demurrer is made pursuant to Code of Civil Procedure sections 430.10 "
            "and 430.30, and is based upon this Notice, the attached Memorandum of Points "
            "and Authorities, the concurrently filed Declaration of ",
            {"text": ATTY_NAME, "text_plain": True},
            " regarding the parties' meet-and-confer efforts, the pleadings and records on "
            "file in this action, and such further evidence and argument as may be presented "
            "at the hearing.",
        ],
    )

    add_body_paragraph(
        doc,
        [
            "The District demurs to the Complaint on the following grounds:",
        ],
    )

    # Enumerated grounds
    for ground_text in [
        (
            "The First Cause of Action (Account Stated) fails to state facts sufficient to "
            "constitute a cause of action against the District.  (Code Civ. Proc., "
            "§ 430.10, subd. (e).)"
        ),
        (
            "The Second Cause of Action (Declaratory Relief) fails to state facts sufficient "
            "to constitute a cause of action against the District and is uncertain.  "
            "(Code Civ. Proc., § 430.10, subds. (e), (f).)"
        ),
        (
            "The Third Cause of Action (Equitable Estoppel) fails to state facts sufficient "
            "to constitute a cause of action against the District.  (Code Civ. Proc., "
            "§ 430.10, subd. (e).)"
        ),
    ]:
        add_body_paragraph(doc, [ground_text])

    add_body_paragraph(
        doc,
        [
            "This Demurrer is supported by the Memorandum of Points and Authorities set "
            "forth below.",
        ],
    )


# ---------------------------------------------------------------------------
# Because some segments above contain a trick (dict without formatting), fix
# add_body_paragraph to accept dicts with only "text" as plain text.
# ---------------------------------------------------------------------------


# Re-define with handling for dict-without-styling-flags — easier than carrying
# a second helper. (Python re-defines silently; the new function replaces the
# earlier one.)


def add_body_paragraph(doc, segments):  # noqa: F811
    p = doc.add_paragraph()
    _apply_body_format(p)
    for seg in segments:
        if isinstance(seg, str):
            run = p.add_run(seg)
            _style_run(run)
        else:
            run = p.add_run(seg["text"])
            _style_run(
                run,
                italic=seg.get("italic", False),
                bold=seg.get("bold", False),
                underline=seg.get("underline", False),
            )
    return p


# ---------------------------------------------------------------------------
# Memorandum
# ---------------------------------------------------------------------------


def build_memorandum(doc):
    add_heading(doc, "MEMORANDUM OF POINTS AND AUTHORITIES")
    add_blank(doc)

    # ---- I. INTRODUCTION ----
    add_heading(doc, "I.")
    add_heading(doc, "INTRODUCTION")
    add_blank(doc)

    add_body_paragraph(
        doc,
        [
            "This is a construction-contract dispute arising out of a publicly bid school "
            "reroofing project.  Plaintiff ",
            {"text": PLAINTIFF_FULL, "italic": False},
            " seeks to recover $187,000 in so-called “punch-list completion” and "
            "warranty-extension charges that the District's governing Board never "
            "approved, that no authorized change order ever documented, and that the "
            "contract itself never contemplated.  The Complaint frames three causes of "
            "action against the District:  account stated, declaratory relief, and "
            "equitable estoppel.  None states a viable claim.",
        ],
    )

    add_body_paragraph(
        doc,
        [
            "Each challenged cause of action founders on the same bedrock principle:  a "
            "public school district contracts only through the means the Legislature has "
            "prescribed, and neither an inspector's informal acknowledgment, a "
            "post-contractual “implied” warranty theory, nor the doctrine of "
            "equitable estoppel may be deployed to circumvent those requirements.  The "
            "Demurrer should be sustained without leave to amend.",
        ],
    )

    add_blank(doc)

    # ---- II. STATEMENT OF FACTS ----
    add_heading(doc, "II.")
    add_heading(doc, "STATEMENT OF FACTS AND RELEVANT ALLEGATIONS")
    add_blank(doc)

    add_body_paragraph(
        doc,
        [
            "On or about April 3, 2024, following a competitive public bid process "
            "conducted pursuant to Public Contract Code section 20111, the District "
            "awarded Plaintiff a contract of approximately $1,600,000 for the full "
            "tear-off and replacement of a 32,000-square-foot built-up roof system over "
            "the gymnasium at Sierra Vista High School (the “Project”).  (Compl., "
            "¶¶ 8–9.)  The written agreement was executed by the "
            "Superintendent pursuant to a written delegation of contracting authority "
            "under Education Code section 17605, and was ratified by the District's "
            "governing Board by Resolution No. 2023–24-41 on April 23, 2024.  "
            "(Compl., ¶ 10.)",
        ],
    )

    add_body_paragraph(
        doc,
        [
            "The written contract contains an integration clause, a two-year workmanship "
            "warranty running from final acceptance, and a change-order article providing "
            "that “[n]o change in the Work shall be undertaken, and no additional "
            "compensation shall be due the Contractor, absent a written change order "
            "executed by the District's authorized representative prior to performance of "
            "the changed Work.”  (Compl., ¶ 12 & Exh. A [Contract], § 7.3.)  "
            "The contract does not incorporate any manufacturer's material warranty, does "
            "not contain a five-year warranty term, and does not condition the District's "
            "payment obligation on future manufacturer performance.  (",
            {"text": "Id.", "italic": True},
            ", §§ 7.3, 9.1.)",
        ],
    )

    add_body_paragraph(
        doc,
        [
            "Plaintiff alleges that during pre-bid discussions with District personnel it "
            "“represented and was understood to have bound the District to” an "
            "additional five-year material warranty provided by the membrane manufacturer, "
            "which Plaintiff contends became an “implied term” of the written "
            "contract that followed.  (Compl., ¶¶ 17–20.)  Plaintiff "
            "further alleges that at substantial completion the District's project "
            "inspector — an outside consultant retained by the District's "
            "construction-management firm — issued a punch-list identifying "
            "thirty-seven items of incomplete or defective work, and orally stated during "
            "walk-throughs that “the punch-list items would be paid as extras.”  "
            "(Compl., ¶¶ 24–27.)  No written change order was ever issued.  "
            "(Compl., ¶ 28.)",
        ],
    )

    add_body_paragraph(
        doc,
        [
            "After completing its work, Plaintiff submitted a final invoice including "
            "$187,000 in additional charges characterized as “punch-list completion” "
            "and “warranty-extension” amounts.  (Compl., ¶ 30.)  Plaintiff "
            "alleges that the District's inspector “acknowledged the amounts owed in "
            "writing” by initialing a line item on the invoice and by transmitting an "
            "email stating that the “figures look correct based on my walkthrough.”  "
            "(Compl., ¶¶ 31–32.)  The District's governing Board has never "
            "approved, ratified, or authorized payment of those additional amounts.  "
            "(Compl., ¶ 34.)  The Complaint quotes the following acknowledgment "
            "from the inspector's punch-list walk-through report as the cornerstone of its "
            "account-stated claim:",
        ],
    )

    add_block_quote(
        doc,
        (
            "I have completed my walk-through of the gymnasium reroofing project and have "
            "reviewed Granite Peak Builders' final invoice dated July 18, 2025.  The "
            "figures look correct based on my walkthrough, and the punch-list items "
            "identified in my May 12, 2025 report should in my view be treated as extras "
            "and paid accordingly.  I will forward this confirmation to the District's "
            "business office."
        ),
        ["(Compl., ¶ 32.)"],
    )

    add_body_paragraph(
        doc,
        [
            "The Complaint pleads three causes of action against the District:  (1) account "
            "stated, premised on the inspector's alleged acknowledgment; (2) declaratory "
            "relief, seeking a judicial declaration that the contract's two-year "
            "workmanship warranty implicitly incorporates an additional five-year "
            "manufacturer's material warranty; and (3) equitable estoppel, seeking to "
            "preclude the District from relying on the written change-order requirement.  "
            "(Compl., ¶¶ 40–67.)  All three causes of action are defective "
            "on the face of the pleading.",
        ],
    )

    add_blank(doc)

    # ---- III. LEGAL STANDARD ----
    add_heading(doc, "III.")
    add_heading(doc, "LEGAL STANDARD ON DEMURRER")
    add_blank(doc)

    add_body_paragraph(
        doc,
        [
            "A demurrer tests the legal sufficiency of the complaint.  On demurrer, the "
            "Court accepts as true all material factual allegations properly pleaded, "
            "together with those facts that may be implied or inferred from the express "
            "allegations.  (",
            {"text": "Blank v. Kirwan", "italic": True},
            " (1985) 39 Cal.3d 311, 318.)  The Court does not, however, “accept "
            "contentions, deductions, or conclusions of fact or law” dressed up as "
            "allegations of ultimate fact.  (",
            {"text": "Zelig v. County of Los Angeles", "italic": True},
            " (2002) 27 Cal.4th 1112, 1126.)  A demurrer reaches only the face of the "
            "complaint and those matters of which the Court may take judicial notice; "
            "evidentiary issues are not at play.  (",
            {"text": "Aubry v. Tri-City Hospital Dist.", "italic": True},
            " (1992) 2 Cal.4th 962, 967.)",
        ],
    )

    add_body_paragraph(
        doc,
        [
            "Where a complaint fails to state a cause of action, leave to amend should be "
            "denied unless the plaintiff shows a reasonable possibility that the defect can "
            "be cured.  The burden of demonstrating amendability rests squarely on the "
            "plaintiff.  (",
            {"text": "Schifando v. City of Los Angeles", "italic": True},
            " (2003) 31 Cal.4th 1074, 1081.)  That burden is not discharged by rote "
            "assertion; the plaintiff must identify how the pleading defect can be cured "
            "and what new facts, if any, would be added.  (",
            {"text": "Id.", "italic": True},
            " at p. 1081.)",
        ],
    )

    add_body_paragraph(
        doc,
        [
            "Those standards are particularly important where, as here, the defendant is "
            "a public school district.  California law has long recognized that public "
            "entities contract only in the manner prescribed by statute, and strict "
            "compliance with those statutory requirements is a jurisdictional condition "
            "precedent to any enforceable contractual obligation.  (",
            {"text": "G.L. Mezzetta, Inc. v. City of American Canyon", "italic": True},
            " (2000) 78 Cal.App.4th 1087, 1092; ",
            {"text": "First Street Plaza Partners v. City of Los Angeles", "italic": True},
            " (1998) 65 Cal.App.4th 650, 669.)",
        ],
    )

    add_blank(doc)

    # ---- IV. ARGUMENT ----
    add_heading(doc, "IV.")
    add_heading(doc, "ARGUMENT")
    add_blank(doc)

    # --- A. Account stated ---
    add_heading(
        doc,
        "A.  THE FIRST CAUSE OF ACTION FOR ACCOUNT STATED FAILS BECAUSE A PROJECT "
        "INSPECTOR LACKS AUTHORITY TO ACKNOWLEDGE A DEBT ON BEHALF OF THE DISTRICT.",
        center=False,
    )
    add_blank(doc)

    add_body_paragraph(
        doc,
        [
            "An account stated requires previous transactions establishing a "
            "debtor-creditor relationship, an express or implied agreement on the amount "
            "due, and an express or implied promise by the debtor to pay.  The "
            "indispensable element is an acknowledgment of the debt by the party to be "
            "charged — here, the District.  Plaintiff's theory that a project "
            "inspector acknowledged the claimed $187,000 obligation is a legal "
            "impossibility.",
        ],
    )

    add_body_paragraph(
        doc,
        [
            "A California public school district contracts only as authorized by the "
            "Education Code and the Public Contract Code.  Education Code section 17604 "
            "vests contracting authority in the governing board and requires that any "
            "contract entered in advance of board approval be ratified by formal board "
            "action before it becomes valid and binding.  Education Code section 17605 "
            "permits the board to delegate contracting authority to the superintendent "
            "(or to such persons as the superintendent designates in writing), but the "
            "delegation must be in writing and must specify the nature and scope of the "
            "authority conferred.  The Legislature's evident intent is that a school "
            "district's contractual obligations be traceable to an identifiable action of "
            "the governing board or an authorized delegee acting within the scope of a "
            "written delegation.",
        ],
    )

    add_body_paragraph(
        doc,
        [
            "Public Contract Code section 20118.4 reinforces the same principle in the "
            "context of change orders on school-district public works contracts, "
            "requiring that changes in scope or price be reduced to writing and executed "
            "by an authorized district representative.  (Pub. Cont. Code, section "
            # ^^ LS-CITE-02 corruption: `§` → `section` inside a parenthetical
            # cite (should be `§` inside parens per LS Manual).
            "20118.4, subd. (a).)  California courts enforce those written change-order "
            "requirements with full rigor.  (",
            {"text": "P&D Consultants, Inc. v. City of Carlsbad", "italic": True},
            " (2010) 190 Cal.App.4th 1332, 1341.)",
        ],
    )

    add_body_paragraph(
        doc,
        [
            "The Complaint identifies the purported acknowledgor as the “District's "
            "project inspector” — an outside consultant retained by the District's "
            "construction-management firm to observe and report on the progress of the "
            "Work.  (Compl., ¶¶ 24, 31.)  Plaintiff does not allege that the "
            "inspector was a member of the District's governing Board, that he held a "
            "written delegation of contracting authority under Education Code section "
            "17605, or that he was ever designated as the District's authorized "
            "representative for purposes of executing change orders under Public Contract "
            "Code section 20118.4.  Plaintiff's own pleading describes the inspector's "
            "role in terms categorically inconsistent with any such authority.  (Compl., "
            "¶ 24 [“retained . . . to observe and report”].)",
        ],
    )

    add_body_paragraph(
        doc,
        [
            "California law is unambiguous that a contract with a public entity entered "
            "in violation of statutory requirements is void, and the doctrines of quantum "
            "meruit and account stated cannot be used to impose a payment obligation the "
            "statutes forbid.  (",
            {"text": "Miller v. McKinnon", "italic": True},
            " (1942) 20 Cal.2d 83, 88–89; ",
            {"text": "Reams v. Cooley", "italic": True},
            # LS-SP-02 corruption: single space between citation sentence
            # "(Miller...; Reams...)" and next body sentence "Contractors..."
            # (was two spaces). Mid-paragraph typo-style miss.
            " (1915) 171 Cal. 150, 154.) Contractors who deal with a public agency are "
            "charged with knowledge of the contracting statutes and bear the risk of "
            "noncompliance.  (",
            {"text": "Katsura v. City of San Buenaventura", "italic": True},
            " (2007) 155 Cal.App.4th 104, 109.)",
        ],
    )

    add_body_paragraph(
        doc,
        [
            "Those principles are dispositive.  If the governing Board never approved "
            "the claimed $187,000 obligation, and if no authorized representative ever "
            "executed a written change order documenting the claimed scope additions, "
            "then the District is not bound to pay those amounts — regardless of "
            "what a project inspector wrote on an invoice line or said during a "
            "walk-through.  An account stated cannot create a debt the underlying law "
            "prohibits.  (",
            {"text": "G.L. Mezzetta, Inc.", "italic": True},
            ", 78 Cal.App.4th at p. 1092; ",
            {"text": "First Street Plaza Partners", "italic": True},
            ", 65 Cal.App.4th at p. 669.)  To hold otherwise would allow the "
            "account-stated doctrine to become an end run around Education Code sections "
            "17604 and 17605 and Public Contract Code section 20118.4 ",
            {"text": "et seq.", "italic": False},
            "  The statutes foreclose that route.",
        ],
    )

    add_body_paragraph(
        doc,
        [
            "The Supreme Court's decision in ",
            {"text": "Amelco Electric v. City of Thousand Oaks", "italic": True},
            " (2002) 27 Cal.4th 228, 234, 242–243 confirms the broader rule:  where "
            "a contractor seeks to recover amounts beyond the contract sum in "
            "circumstances the contract and governing statutes do not authorize, the "
            "claim fails as a matter of law.  The First Cause of Action should "
            "therefore be sustained without leave to amend; Plaintiff cannot cure the "
            "defect by alleging additional inspector statements, because the problem is "
            "structural rather than evidentiary.  (",
            {"text": "Schifando", "italic": True},
            ", 31 Cal.4th at p. 1081.)",
        ],
    )

    add_blank(doc)

    # --- B. Declaratory relief ---
    add_heading(
        doc,
        "B.  THE SECOND CAUSE OF ACTION FOR DECLARATORY RELIEF FAILS BECAUSE THE "
        "IMPLIED COVENANT CANNOT REWRITE THE EXPRESS TERMS OF THE PARTIES' WRITTEN "
        "CONTRACT.",
        center=False,
    )
    add_blank(doc)

    add_body_paragraph(
        doc,
        [
            "The second cause of action seeks a judicial declaration that the contract's "
            "two-year workmanship warranty “implicitly incorporates” a five-year "
            "manufacturer's material warranty that Plaintiff allegedly bound the "
            "District to accept during pre-bid discussions.  (Compl., ¶¶ "
            "50–55.)  The claim invites the Court to rewrite the parties' written "
            "contract through the back door of an implied-covenant theory; California "
            "law does not permit that.",
        ],
    )

    add_body_paragraph(
        doc,
        [
            "The essential elements of a contract — parties capable of contracting, "
            "mutual consent, a lawful object, and sufficient cause or consideration — "
            "are prescribed by Civil Code section 1550, and the requirement of mutual "
            "consent is further defined in Civil Code section 1565.  Where the parties "
            "have reduced their agreement to a fully integrated writing, California law "
            "treats the writing as the complete and exclusive statement of their terms "
            "and declines to look behind it to pre-contractual discussions absent grounds "
            "recognized by the parol evidence rule.",
        ],
    )

    add_body_paragraph(
        doc,
        [
            "The covenant of good faith and fair dealing — on which Plaintiff's "
            "theory ultimately depends — does not provide an escape valve.  The "
            "implied covenant cannot be endowed with an existence independent of its "
            "contractual underpinnings and cannot impose substantive duties or limits on "
            "the contracting parties beyond those incorporated in the specific terms of "
            "their agreement.  (",
            {"text": "Carma Developers (Cal.), Inc. v. Marathon Development California, Inc.", "italic": True},
            " (1992) 2 Cal.4th 342, 374.)  The Supreme Court has since reiterated that "
            "the implied covenant is limited to assuring compliance with the express "
            "terms of the contract and cannot be extended to create substantive terms "
            "inconsistent with the written agreement.  (",
            {"text": "Guz v. Bechtel National, Inc.", "italic": True},
            " (2000) 24 Cal.4th 317, 349–350.)",
        ],
    )

    add_body_paragraph(
        doc,
        [
            "Plaintiff's theory runs headlong into those authorities.  The written "
            "contract provides a two-year workmanship warranty — not seven years, "
            "not five years plus two, not any period derived from extrinsic manufacturer "
            "documentation.  (Compl., Exh. A [Contract], § 9.1.)  The contract "
            "does not incorporate the manufacturer's warranty by reference, and "
            "Plaintiff does not allege that it does.  To hold that the two-year warranty "
            "“implicitly incorporates” a separate five-year warranty would not "
            "give effect to the contract's terms; it would replace them.  (",
            {"text": "Carma", "italic": True},
            ", 2 Cal.4th at p. 374; ",
            {"text": "Guz", "italic": True},
            ", 24 Cal.4th at p. 349.)",
        ],
    )

    add_body_paragraph(
        doc,
        [
            "The public-entity dimension of this case deepens the problem.  Plaintiff's "
            "theory is not merely that the parties' course of dealing added a warranty "
            "term; it is that Plaintiff's pre-bid representations to District personnel "
            "became binding on the District notwithstanding the integrated writing.  As "
            "applied to a public school district, that theory collides with the "
            "contracting statutes:  the District may bind itself to a warranty obligation "
            "only through the means prescribed by Education Code sections 17604 and "
            # LS-SP-02 corruption: single space between sentences
            # "17605. Pre-bid" (was two spaces). Mid-paragraph, deep in the
            # declaratory-relief argument.
            "17605. Pre-bid representations by prospective bidders to district staff are "
            "not among those means.  (",
            {"text": "First Street Plaza Partners", "italic": True},
            ", 65 Cal.App.4th at p. 669; ",
            {"text": "G.L. Mezzetta, Inc.", "italic": True},
            ", 78 Cal.App.4th at p. 1092.)",
        ],
    )

    add_body_paragraph(
        doc,
        [
            "Nor does Government Code section 53060 — the provision generally "
            "authorizing public agencies to contract for special services — assist "
            "Plaintiff; that section does not displace the specific statutory framework "
            "governing school-district construction contracts and does not authorize the "
            "agency to be bound by unwritten representations exchanged before a formal "
            "contract is executed.  (Gov. Code, § 53060.)  The claim is also "
            "uncertain, because Plaintiff does not identify the manufacturer whose "
            "warranty is allegedly incorporated, does not attach or quote the text of "
            "that warranty, and does not specify which “seam failures” fall within "
            "its scope.  A declaratory-relief cause of action that asks the Court to "
            "decree the terms of a warranty neither identified nor pleaded is uncertain "
            "within the meaning of Code of Civil Procedure section 430.10, subdivision "
            "(f).  The Second Cause of Action should be sustained without leave to amend.",
        ],
    )

    add_blank(doc)

    # --- C. Equitable estoppel ---
    add_heading(
        doc,
        "C.  THE THIRD CAUSE OF ACTION FOR EQUITABLE ESTOPPEL FAILS BECAUSE ESTOPPEL "
        "MAY NOT BE USED TO DEFEAT THE STATUTORY BIDDING AND CHANGE-ORDER FRAMEWORK.",
        center=False,
    )
    add_blank(doc)

    add_body_paragraph(
        doc,
        [
            "Plaintiff's third cause of action seeks to preclude the District from "
            "invoking the written change-order requirement and the bidding and "
            "contracting framework, based on the inspector's alleged oral statements that "
            "“the punch-list items would be paid as extras.”  (Compl., "
            "¶¶ 60–67.)  Equitable estoppel against a public entity is "
            "available, if at all, only in the rarest of circumstances, and it is "
            "categorically unavailable where, as here, it would nullify a statutory "
            "policy adopted for the protection of the public.",
        ],
    )

    add_body_paragraph(
        doc,
        [
            "The governing decision is ",
            {"text": "City of Long Beach v. Mansell", "italic": True},
            " (1970) 3 Cal.3d 462, 493, 496–497, in which the Supreme Court "
            "recognized that equitable estoppel may, in unusual circumstances, be "
            "invoked against a governmental body, but held that estoppel cannot be "
            "applied where its effect would be to nullify a strong rule of policy "
            "adopted for the benefit of the public.  That two-part framework has "
            "governed California public-entity estoppel cases for more than a half "
            "century.",
        ],
    )

    add_body_paragraph(
        doc,
        [
            "The Supreme Court reaffirmed that rule in ",
            {"text": "Lentz v. McMahon", "italic": True},
            " (1989) 49 Cal.3d 393, 399, holding that estoppel may not be invoked "
            "against a public entity where doing so would nullify a strong rule of "
            "policy adopted for the benefit of the public.  And in ",
            {"text": "Janis v. California State Lottery Com.", "italic": True},
            " (1998) 68 Cal.App.4th 824, 829–830, the Court of Appeal reiterated "
            "that estoppel against a public entity is rare and limited, reserved for the "
            "exceptional case in which its application would not compromise an "
            "overriding statutory objective.",
        ],
    )

    add_body_paragraph(
        doc,
        [
            "The statutory scheme here embodies precisely the kind of strong public "
            "policy the ",
            {"text": "Mansell", "italic": True},
            " line of cases protects.  The public-bidding statutes ensure that "
            "school-district funds are expended through an open and competitive process, "
            "not through private arrangements with favored contractors.  (Pub. Cont. "
            "Code, § 20111.)  The written change-order requirement of Public "
            "Contract Code section 20118.4 exists so that modifications to publicly bid "
            "contracts are documented, traceable, and subject to board review.  (Pub. "
            "Cont. Code, § 20118.4, subd. (b).)  And the delegation and ratification "
            "provisions of Education Code sections 17604 and 17605 ensure that the "
            "District's contractual commitments are made by, or on behalf of, the elected "
            "body the statutes designate as the contracting authority.",
        ],
    )

    add_body_paragraph(
        doc,
        [
            "Applying estoppel to bar the District from invoking those requirements "
            "would disable every one of them.  If an inspector's oral statement can "
            "preclude a school district from insisting on a written change order, the "
            "change-order statute becomes a dead letter; if pre-bid representations can "
            # LS-CAP-02 corruption: "the District" (defined party-substitute
            # for Sierra Vista Unified School District) lowercased to
            # "the district" deep in the estoppel argument.
            "bind the district despite the absence of board action, the Education Code's "
            "delegation and ratification provisions fail in their purpose.  That is "
            "exactly the outcome ",
            {"text": "Mansell", "italic": True},
            ", ",
            {"text": "Lentz", "italic": True},
            ", and ",
            {"text": "Janis", "italic": True},
            " foreclose.  (",
            {"text": "Mansell", "italic": True},
            ", 3 Cal.3d at p. 496; ",
            {"text": "Lentz", "italic": True},
            ", 49 Cal.3d at p. 399; ",
            {"text": "Janis", "italic": True},
            ", 68 Cal.App.4th at p. 830.)",
        ],
    )

    add_body_paragraph(
        doc,
        [
            "The Court of Appeal's decision in ",
            {"text": "Air Quality Products, Inc. v. State of California", "italic": True},
            # LS-CITE-HAL corruption: first-page digits transposed from
            # "340" to "304" (real cite is 96 Cal.App.3d 340). The pin-page
            # range "348–349" stays; only the first-page digit transposition
            # breaks the CourtListener lookup.
            " (1979) 96 Cal.App.3d 304, 348–349 provides an independent ground for "
            "the same conclusion:  estoppel cannot be invoked against a public agency "
            "where, as here, the conduct relied on exceeded the agency's statutory "
            "authority.  The inspector had no power to execute change orders and no power "
            "to bind the District to payment of extras.  (Ed. Code, §§ 17604, "
            "17605; Pub. Cont. Code, § 20118.4.)  When a contractor bids on a "
            "school-district public works project, it accepts the "
            "contracting rules the Legislature has prescribed; those rules are the "
            "contractor's responsibility to know and follow, and noncompliance is the "
            "contractor's risk, not the district's.  (",
            {"text": "Miller", "italic": True},
            ", 20 Cal.2d at p. 89; ",
            {"text": "Katsura", "italic": True},
            ", 155 Cal.App.4th at p. 109.)  Equitable estoppel is not a mechanism for "
            "reallocating that risk to the public.  The Third Cause of Action should be "
            "sustained without leave to amend.",
        ],
    )

    add_blank(doc)

    # ---- V. CONCLUSION ----
    add_heading(doc, "V.")
    add_heading(doc, "CONCLUSION")
    add_blank(doc)

    add_body_paragraph(
        doc,
        [
            "For the reasons set forth above, each of the three challenged causes of "
            "action — account stated, declaratory relief, and equitable estoppel "
            "— fails as a matter of law on the face of the Complaint, and no "
            "amendment can cure the underlying defects.  The District respectfully "
            "requests that the Court sustain the Demurrer in its entirety, without leave "
            "to amend, and enter judgment in favor of the District on the first, second, "
            "and third causes of action.",
        ],
    )

    add_blank(doc)

    # Signature block
    sig_lines = [
        f"Dated:  {SIGNATURE_DATE}",
        f"{FIRM_NAME}",
        "",
        "By:  ____________________________________",
        f"     {ATTY_NAME}",
        f"     {ATTY_2_NAME}",
        f"     Attorneys for Defendant",
        f"     {DISTRICT_FULL.upper()}",
    ]
    for line in sig_lines:
        p = doc.add_paragraph()
        _apply_plain_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        if line:
            run = p.add_run(line)
            _style_run(run)


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------


def main():
    doc = Document()

    # Page setup — 1" margins all around (standard CA trial court filing).
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Make sure the default style is TNR 12 — belt-and-suspenders, since we
    # also set every run individually.
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = FONT_SIZE
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), FONT_NAME)

    build_caption(doc)
    build_notice_and_demurrer(doc)
    build_memorandum(doc)

    out = Path(__file__).resolve().parent.parent / "realistic-mixed.docx"
    doc.save(str(out))
    print(f"Wrote {out} ({out.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
