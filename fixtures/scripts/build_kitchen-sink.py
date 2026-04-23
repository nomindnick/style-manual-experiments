"""Build the kitchen-sink-violations fixture.

Produces `fixtures/kitchen-sink-violations.docx` — a Phase-0.5 corrupted
fixture derived from the Eastvale clean Wave-A draft (now superseded;
audit trail in `fixtures/kitchen-sink-violations.SOURCE_NOTES.md` and
`.SOURCE_VIOLATIONS.md`). This document deliberately contains ~17-18 LS
Style Manual violations spread across the Phase-1 rules plus neighbor
coverage. Ground-truth violations are logged alongside in
`fixtures/kitchen-sink-violations.violations.json`.

Each injected violation is marked with a `# CORRUPT(<rule_id>)` comment
so a reader can find the planted-error edits at a glance.

Run from repo root:

    .venv/bin/python fixtures/scripts/build_kitchen-sink.py

See `fixtures/CORRUPT_BRIEF.md` for the full corruption spec.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


# --------------------------------------------------------------------------- #
# Formatting constants
# --------------------------------------------------------------------------- #

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(12)
BODY_LINE_SPACING = Pt(24)  # "exactly 24 pt"
BODY_FIRST_LINE_INDENT = Inches(0.5)
BLOCK_INDENT = Inches(0.5)

# Two spaces between sentences (LS-SP-02).  Used throughout the body text.
SS = "  "

OUTPUT_PATH = (
    Path(__file__).resolve().parent.parent / "kitchen-sink-violations.docx"
)


# --------------------------------------------------------------------------- #
# Helpers
# --------------------------------------------------------------------------- #


def _apply_run_font(run, *, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    # python-docx sometimes sets only the "ascii" font; set east-asian too so
    # Word doesn't fall back on Calibri for ambiguous characters.
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rFonts.set(qn(f"w:{attr}"), FONT_NAME)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _apply_superscript(run) -> None:
    """CORRUPT(LS-CITE-08b) helper — set vertAlign=superscript on a run."""
    rPr = run._element.get_or_add_rPr()
    vAlign = rPr.find(qn("w:vertAlign"))
    if vAlign is None:
        vAlign = OxmlElement("w:vertAlign")
        rPr.append(vAlign)
    vAlign.set(qn("w:val"), "superscript")


def _set_body_format(paragraph, *, first_line_indent: bool = True) -> None:
    pf = paragraph.paragraph_format
    pf.line_spacing = BODY_LINE_SPACING
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if first_line_indent:
        pf.first_line_indent = BODY_FIRST_LINE_INDENT


def add_body_paragraph(doc, segments, *, first_line_indent: bool = True, align=None):
    """Add a body paragraph composed of (text, style_flags) segments.

    `segments` is a list of either a plain string or a tuple
    ``(text, {"italic": True/False, "bold": True/False, "superscript": True})``.
    """
    p = doc.add_paragraph()
    _set_body_format(p, first_line_indent=first_line_indent)
    if align is not None:
        p.alignment = align
    for seg in segments:
        if isinstance(seg, str):
            text, flags = seg, {}
        else:
            text, flags = seg
        run = p.add_run(text)
        _apply_run_font(
            run,
            bold=flags.get("bold"),
            italic=flags.get("italic"),
        )
        if flags.get("superscript"):
            _apply_superscript(run)
    return p


def add_heading(doc, text, *, center: bool = False) -> None:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = BODY_LINE_SPACING
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = None
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    _apply_run_font(run, bold=True)


def add_block_quote(doc, text, citation_text) -> None:
    """Block quote indented 0.5" both sides, single-spaced, no first-line indent."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = BLOCK_INDENT
    pf.right_indent = BLOCK_INDENT
    pf.first_line_indent = None
    pf.line_spacing = 1.0  # single-spaced
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(12)
    run = p.add_run(text)
    _apply_run_font(run)
    # Citation on its own line, at left margin, normal indentation.
    cp = doc.add_paragraph()
    _set_body_format(cp, first_line_indent=True)
    crun = cp.add_run(citation_text)
    _apply_run_font(crun)


def add_blank_line(doc) -> None:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = BODY_LINE_SPACING
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


# --------------------------------------------------------------------------- #
# Caption block
# --------------------------------------------------------------------------- #


def build_caption(doc) -> None:
    # Attorneys-of-record block (left side, above caption)
    atty_lines = [
        "LOZANO SMITH",
        "Margaret A. Delacroix, Esq. (SBN 289431)",
        "Jonathan R. Pemberton, Esq. (SBN 312778)",
        "One Capitol Mall, Suite 640",
        "Sacramento, California 95814",
        "Telephone:  (916) 329-7433",
        "Facsimile:  (916) 329-9050",
        "Email:  mdelacroix@lozanosmith.com",
        "",
        "Attorneys for Defendant",
        "EASTVALE JOINT UNION HIGH SCHOOL DISTRICT",
    ]
    for line in atty_lines:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = BODY_LINE_SPACING
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.first_line_indent = None
        run = p.add_run(line)
        _apply_run_font(run, bold=(line == "LOZANO SMITH"))

    add_blank_line(doc)

    # Court name — centered, bold
    for line in (
        "SUPERIOR COURT OF THE STATE OF CALIFORNIA",
        "FOR THE COUNTY OF SAN BERNARDINO",
    ):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = BODY_LINE_SPACING
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.first_line_indent = None
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        _apply_run_font(run, bold=True)

    add_blank_line(doc)

    # Caption table: left = parties; right = case info
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(3.3)
    table.columns[1].width = Inches(3.2)

    left_cell = table.cell(0, 0)
    right_cell = table.cell(0, 1)
    left_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    right_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Left cell: plaintiff v. defendant
    left_lines = [
        ("RIVERBEND MECHANICAL SERVICES, LLC,", False),
        ("a California limited liability company,", False),
        ("", False),
        ("                      Plaintiff,", False),
        ("", False),
        ("          v.", False),
        ("", False),
        ("EASTVALE JOINT UNION HIGH SCHOOL", False),
        ("DISTRICT; and DOES 1 through 25,", False),
        ("inclusive,", False),
        ("", False),
        ("                      Defendants.", False),
    ]
    # Clear initial empty paragraph in left cell
    left_cell.paragraphs[0].text = ""
    first = True
    for text, bold in left_lines:
        if first:
            p = left_cell.paragraphs[0]
            first = False
        else:
            p = left_cell.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = BODY_LINE_SPACING
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.first_line_indent = None
        run = p.add_run(text)
        _apply_run_font(run, bold=bold)

    # Right cell: case number, title, hearing info
    right_lines = [
        ("Case No. CIVSB 2532109", False),
        ("", False),
        ("DEFENDANT EASTVALE JOINT UNION HIGH", True),
        ("SCHOOL DISTRICT'S DEMURRER TO", True),
        ("PLAINTIFF'S COMPLAINT; MEMORANDUM", True),
        ("OF POINTS AND AUTHORITIES IN", True),
        ("SUPPORT THEREOF", True),
        ("", False),
        ("[Filed concurrently with Request for Judicial", False),
        ("Notice; Declaration of Margaret A. Delacroix", False),
        ("re: Meet and Confer; and [Proposed] Order]", False),
        ("", False),
        ("Hearing Date:  July 14, 2026", False),
        ("Hearing Time:  8:30 a.m.", False),
        ("Department:  S-26", False),
        ("Judge:  Hon. Teresa M. Halvorsen", False),
        ("", False),
        ("Action Filed:  February 11, 2026", False),
        ("Trial Date:  Not Set", False),
    ]
    right_cell.paragraphs[0].text = ""
    first = True
    for text, bold in right_lines:
        if first:
            p = right_cell.paragraphs[0]
            first = False
        else:
            p = right_cell.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = BODY_LINE_SPACING
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.first_line_indent = None
        run = p.add_run(text)
        _apply_run_font(run, bold=bold)


# --------------------------------------------------------------------------- #
# Body sections
# --------------------------------------------------------------------------- #


def build_notice(doc) -> None:
    add_blank_line(doc)
    add_heading(doc, "NOTICE OF DEMURRER AND DEMURRER", center=True)
    add_blank_line(doc)
    add_body_paragraph(
        doc,
        [
            "TO PLAINTIFF RIVERBEND MECHANICAL SERVICES, LLC AND ITS ATTORNEYS OF RECORD:",
        ],
        first_line_indent=False,
    )
    add_blank_line(doc)
    add_body_paragraph(
        doc,
        [
            "PLEASE TAKE NOTICE that on July 14, 2026, at 8:30 a.m., or as soon "
            "thereafter as the matter may be heard, in Department S-26 of the above-"
            "entitled Court, located at 247 West Third Street, San Bernardino, "
            "California 92415, defendant Eastvale Joint Union High School District "
            "(“District”) will, and hereby does, demur to the Complaint of "
            "plaintiff Riverbend Mechanical Services, LLC (“Plaintiff”) filed "
            "on February 11, 2026."
        ],
    )
    add_body_paragraph(
        doc,
        [
            "The District demurs to the Complaint, and to the first, second, and "
            "third causes of action set forth therein, on the following grounds:"
        ],
    )
    # Enumerated grounds — keep as body paragraphs with numeric prefixes.
    grounds = [
        (
            "1.",
            "The first cause of action for breach of the implied covenant of "
            "good faith and fair dealing fails to state facts sufficient to "
            "constitute a cause of action against the District.  (Code Civ. "
            "Proc., § 430.10, subd. (e).)",
        ),
        (
            "2.",
            "The second cause of action for breach of contract fails to state "
            "facts sufficient to constitute a cause of action against the "
            "District and is uncertain.  (Code Civ. Proc., § 430.10, subds. "
            "(e), (f).)",
        ),
        (
            "3.",
            "The third cause of action for unjust enrichment fails to state "
            "facts sufficient to constitute a cause of action against the "
            "District.  (Code Civ. Proc., § 430.10, subd. (e).)",
        ),
    ]
    for num, text in grounds:
        add_body_paragraph(doc, [f"{num}{SS}{text}"])

    add_body_paragraph(
        doc,
        [
            "This Demurrer is based on this Notice; the attached Memorandum of "
            "Points and Authorities; the concurrently-filed Request for Judicial "
            "Notice and Declaration of Margaret A. Delacroix regarding meet-and-"
            "confer efforts under Code of Civil Procedure section 430.41; the "
            "pleadings, papers, and records on file in this action; and any "
            "further evidence and argument the Court may consider at the hearing."
        ],
    )
    add_blank_line(doc)
    add_body_paragraph(
        doc,
        [
            "Dated:  April 22, 2026",
            "".ljust(20),
            ("LOZANO SMITH", {"bold": True}),
        ],
        first_line_indent=False,
    )
    add_blank_line(doc)
    add_body_paragraph(
        doc,
        [
            "".ljust(44),
            "By:  ",
            ("Margaret A. Delacroix", {"bold": True}),
        ],
        first_line_indent=False,
    )
    add_body_paragraph(
        doc,
        [
            "".ljust(44),
            "Margaret A. Delacroix",
        ],
        first_line_indent=False,
    )
    add_body_paragraph(
        doc,
        [
            "".ljust(44),
            "Attorneys for Defendant",
        ],
        first_line_indent=False,
    )
    add_body_paragraph(
        doc,
        [
            "".ljust(44),
            "EASTVALE JOINT UNION HIGH",
        ],
        first_line_indent=False,
    )
    add_body_paragraph(
        doc,
        [
            "".ljust(44),
            "SCHOOL DISTRICT",
        ],
        first_line_indent=False,
    )


def build_memorandum(doc) -> None:
    add_blank_line(doc)
    add_heading(
        doc,
        "MEMORANDUM OF POINTS AND AUTHORITIES",
        center=True,
    )
    add_blank_line(doc)

    # I. INTRODUCTION
    add_heading(doc, "I.  INTRODUCTION")
    # CORRUPT(LS-CAP-02) @ ¶38 — "a California school district" -> "a California school District"
    # (generic reference erroneously capitalized as if it were the defined term)
    add_body_paragraph(
        doc,
        [
            "This action arises out of a public-works contract for the heating, "
            "ventilation, and air-conditioning (“HVAC”) retrofit of the main "
            "classroom building at Eastvale High School.  Plaintiff Riverbend "
            "Mechanical Services, LLC was the low responsible bidder on that "
            "project and was awarded a written contract in the approximate sum of "
            "$2.8 million following competitive bidding.",
            "  The contract — like every contract to which a California school "
            "District is a party — required that any change to scope or price "  # CORRUPT(LS-CAP-02)
            "be memorialized in a written change order executed by an officer "
            "expressly authorized by the District’s governing board.  Plaintiff "
            "now asks this Court to override those statutory and contractual "
            "requirements and to award Plaintiff payment for work that was never "
            "approved by the District’s Board of Trustees and never reduced to "
            "any signed change order.",
        ],
    )
    # CORRUPT(LS-SP-04) @ ¶39 — "premise:  that" -> "premise: that" (one space after colon)
    add_body_paragraph(
        doc,
        [
            "The three causes of action to which the District demurs each rest on "
            "the same flawed premise: that verbal directions from a District "  # CORRUPT(LS-SP-04)
            "maintenance employee can bind the District to pay for scope "
            "additions costing hundreds of thousands of dollars.  That premise is "
            "contrary to settled California law.  The implied covenant of good "
            "faith and fair dealing cannot create obligations the parties did not "
            "write into their contract.  (",
            ("Guz v. Bechtel National, Inc.", {"italic": True}),
            " (2000) 24 Cal.4th 317, 349-350.)  Verbal change orders are not "
            "enforceable against a school district; written change orders, and "
            "board approval of them, are jurisdictional prerequisites.  (Pub. "
            "Cont. Code, § 20118.4; Ed. Code, §§ 17604, 17605.)  And "
            "quantum-meruit or unjust-enrichment recovery is categorically "
            "unavailable against a public entity for work performed outside a "
            "validly-authorized public contract.  (",
            ("Miller v. McKinnon", {"italic": True}),
            " (1942) 20 Cal.2d 83, 88-89.)",
        ],
    )
    # CORRUPT(LS-CAP-02) @ ¶40 — "the District respectfully requests" -> "the district respectfully requests"
    # (defined term erroneously lowercased)
    add_body_paragraph(
        doc,
        [
            "For the reasons set forth below, the district respectfully requests "  # CORRUPT(LS-CAP-02)
            "that the Court sustain this demurrer without leave to amend."
        ],
    )

    # II. STATEMENT OF FACTS
    add_blank_line(doc)
    add_heading(doc, "II.  STATEMENT OF RELEVANT ALLEGATIONS")
    # CORRUPT(LS-SP-02) @ ¶43 — "318.)  The District" -> "318.) The District" (single space between sentences)
    add_body_paragraph(
        doc,
        [
            "For purposes of this demurrer only, the District accepts the "
            "well-pleaded factual allegations of the Complaint as true.  (",
            ("Blank v. Kirwan", {"italic": True}),
            " (1985) 39 Cal.3d 311, 318.) The District does not accept "  # CORRUPT(LS-SP-02)
            "contentions, deductions, or conclusions of law.  (",
            ("Zelig v. County of Los Angeles", {"italic": True}),
            " (2002) 27 Cal.4th 1112, 1126.)  The allegations material to this "
            "demurrer are these.",
        ],
    )
    add_body_paragraph(
        doc,
        [
            "On or about June 4, 2025, the District’s Board of Trustees adopted "
            "Resolution No. 24/25-38 awarding a public-works contract to "
            "Plaintiff, as the low responsible bidder under Public Contract Code "
            "section 20111, for the HVAC retrofit of the main classroom building "
            "at Eastvale High School (the “Project”).  (Compl., ¶¶ "
            "6-8.)  The contract was executed on June 18, 2025, for a stated "
            "contract sum of $2,847,500, with substantial completion required by "
            "August 1, 2026.  (Compl., ¶ 9.)"
        ],
    )
    add_body_paragraph(
        doc,
        [
            "The contract incorporated by reference a set of General Conditions.  "
            "(Compl., ¶ 10 & Ex. A thereto.)  Article 12 of those General "
            "Conditions, titled “Changes in the Work,” is quoted at length "
            "in Plaintiff’s Complaint and provides, in relevant part:"
        ],
    )
    add_block_quote(
        doc,
        "No change in the scope, price, or time of performance of this Contract "
        "shall be effective unless and until such change is memorialized in a "
        "written change order executed by the District’s Superintendent or "
        "other officer expressly authorized by the Board of Trustees to execute "
        "change orders on the District’s behalf, and approved by the Board of "
        "Trustees where the cumulative value of change orders exceeds the "
        "threshold specified in Public Contract Code section 20118.4.  No "
        "verbal direction, field instruction, or course of conduct shall be "
        "deemed a change order or operate to modify this Contract.",
        "(Compl., ¶ 11 & Ex. A, art. 12.2.)",
    )
    add_body_paragraph(
        doc,
        [
            "Plaintiff alleges that, beginning in or about September 2025, the "
            "District’s Director of Maintenance and Operations, Raymond Ortiz, "
            "verbally directed Plaintiff to perform additional scope items not "
            "included in the original contract.  (Compl., ¶¶ 13-15.)  "
            "Specifically, Plaintiff alleges Mr. Ortiz directed Plaintiff to "
            "reroute approximately 1,200 linear feet of supply and return "
            "ductwork to accommodate newly-discovered asbestos-containing "
            "materials, and to install approximately 340 linear feet of "
            "additional condensate-line piping and three supplemental pump "
            "packages.  (Compl., ¶¶ 14, 16-17.)  Plaintiff does not allege "
            "that Mr. Ortiz held any position to which the Board of Trustees had "
            "delegated contracting authority; rather, Plaintiff alleges only "
            "that Mr. Ortiz was “the on-site District representative most "
            "familiar with the Project.”  (Compl., ¶ 15.)"
        ],
    )
    add_body_paragraph(
        doc,
        [
            "Plaintiff alleges that it submitted five proposed written change "
            "orders for this additional work, totaling $487,612, between October "
            "2025 and January 2026.  (Compl., ¶ 19.)  The District’s "
            "Superintendent declined to execute any of the five proposed change "
            "orders, and the Board of Trustees was never presented with, and did "
            "not approve, any of them.  (Compl., ¶¶ 20-22.)  Plaintiff "
            "nevertheless completed the additional work and, on February 2, 2026, "
            "submitted an invoice for the full $487,612, which the District "
            "declined to pay.  (Compl., ¶¶ 23-25.)  Plaintiff filed the "
            "Complaint nine days later.  (Compl., ¶ 1.)"
        ],
    )
    add_body_paragraph(
        doc,
        [
            "The Complaint pleads three causes of action against the District:  "
            "(1) breach of the implied covenant of good faith and fair dealing; "
            "(2) breach of contract premised on the alleged verbal change orders; "
            "and (3) unjust enrichment for the reasonable value of the additional "
            "work.  (Compl., ¶¶ 27-52.)  The District demurs to each."
        ],
    )

    # III. LEGAL STANDARD
    add_blank_line(doc)
    add_heading(doc, "III.  LEGAL STANDARD")
    # CORRUPT(LS-CITE-08b) @ ¶53 — superscript on the "4th" in "2 Cal.4th 962".
    # Split the run around the "4th" token and mark the inner run superscript.
    add_body_paragraph(
        doc,
        [
            "A demurrer tests the legal sufficiency of the complaint.  (",
            ("Aubry v. Tri-City Hospital Dist.", {"italic": True}),
            " (1992) 2 Cal.",
            ("4th", {"superscript": True}),  # CORRUPT(LS-CITE-08b)
            " 962, 966-967.)  In reviewing the sufficiency of a "
            "complaint against a general demurrer, the court treats the demurrer "
            "as admitting all material facts properly pleaded, but does not "
            "assume the truth of contentions, deductions, or conclusions of law.  "
            "(",
            ("Blank", {"italic": True}),
            ", 39 Cal.3d at p. 318.)  The demurrer reaches only the face of the "
            "complaint and matters properly subject to judicial notice.  (",
            ("Aubry", {"italic": True}),
            ", 2 Cal.4th at p. 967.)",
        ],
    )
    add_body_paragraph(
        doc,
        [
            "Where, as here, the plaintiff opposes a demurrer to a particular "
            "cause of action, the plaintiff bears the burden of demonstrating a "
            "reasonable possibility that the defect can be cured by amendment.  "
            "(",
            ("Schifando v. City of Los Angeles", {"italic": True}),
            " (2003) 31 Cal.4th 1074, 1081.)  If the plaintiff cannot meet that "
            "burden, the demurrer should be sustained without leave to amend.  (",
            ("Schifando", {"italic": True}),
            ", 31 Cal.4th at p. 1081.)",
        ],
    )
    # CORRUPT(LS-CITE-03) @ ¶55 — "Under section 430.10, a party may" ->
    # "Under section 430.10, subd. (e), a party may" (outside-parens "subd." should be "subdivision")
    add_body_paragraph(
        doc,
        [
            "The demurrer procedure is governed by Code of Civil Procedure "
            "sections 430.10 et seq.  Under section 430.10, subd. (e), a party "  # CORRUPT(LS-CITE-03)
            "may demur on the ground, among "
            "others, that “[t]he pleading does not state facts sufficient to "
            "constitute a cause of action” or that the pleading is uncertain.  "
            "(Code Civ. Proc., § 430.10, subds. (e), (f).)  The District and "
            "Plaintiff have complied with the meet-and-confer requirements of "
            "Code of Civil Procedure section 430.41, as set forth in the "
            "concurrently-filed Declaration of Margaret A. Delacroix.",
        ],
    )

    # IV. ARGUMENT
    add_blank_line(doc)
    add_heading(doc, "IV.  ARGUMENT")
    add_blank_line(doc)

    # IV.A
    add_heading(
        doc,
        "A.  THE FIRST CAUSE OF ACTION FAILS BECAUSE THE IMPLIED COVENANT CANNOT CREATE OBLIGATIONS NOT IN THE WRITTEN CONTRACT.",
    )
    # CORRUPT(LS-QUOTE-02) @ ¶60 — put semicolon inside quotation marks.
    # Change "...authorized by District personnel" and "refusing..." to
    # "...authorized by District personnel;" and "refusing..." (semicolon inside quote).
    add_body_paragraph(
        doc,
        [
            "Plaintiff’s first cause of action alleges that the District "
            "breached the implied covenant of good faith and fair dealing by "
            "“withholding payment for work authorized by District personnel;” "  # CORRUPT(LS-QUOTE-02)
            "and “refusing to process in good faith the change orders "
            "submitted by Plaintiff.”  (Compl., ¶¶ 28-30.)  The claim "
            "fails as a matter of law because it seeks to impose obligations on "
            "the District that the written contract does not impose — and "
            "indeed that the written contract, and California public-contracting "
            "law, expressly foreclose."
        ],
    )
    add_body_paragraph(
        doc,
        [
            "The implied covenant of good faith and fair dealing is, as its name "
            "suggests, an implied term — it is an interpretive aid that "
            "prevents a contracting party from using the literal terms of a "
            "contract to frustrate the benefits the parties actually bargained "
            "for.  It does not add new substantive terms to an agreement.  Our "
            "Supreme Court has been emphatic on this point:  the implied "
            "covenant has no existence independent of the contractual "
            "relationship on which it operates, and it cannot impose "
            "substantive duties on the parties beyond those incorporated in "
            "the specific terms of their agreement.  (",
            ("Guz v. Bechtel National, Inc.", {"italic": True}),
            " (2000) 24 Cal.4th 317, 349-350.)  The rule operates as a "
            "categorical limit, not a mere gloss on contractual "
            "interpretation.  (",
            ("Id.", {"italic": True}),
            " at p. 350.)",
        ],
    )
    add_body_paragraph(
        doc,
        [
            "The ",
            ("Carma", {"italic": True}),
            " court stated the rule to the same effect.  The implied covenant "
            "is confined to assuring compliance with the express terms of the "
            "contract, and cannot be extended to create obligations the "
            "parties did not themselves agree to undertake.  (",
            ("Carma Developers (Cal.), Inc. v. Marathon Development California, Inc.", {"italic": True}),
            " (1992) 2 Cal.4th 342, 373-374.)  An implied-covenant claim that "
            "is not tethered to a specific express contract term the "
            "defendant frustrated is not a cognizable claim under California "
            "law.  (",
            ("Guz", {"italic": True}),
            ", 24 Cal.4th at p. 350.)"
        ],
    )
    # CORRUPT(LS-CITE-06c) @ ¶63 — remove italics from the trailing "Id." in
    # "(Id. at pp. 349-350.)" — i.e. the Id. run is not italic.
    add_body_paragraph(
        doc,
        [
            "Plaintiff’s first cause of action falls squarely within the ",
            ("Guz/Carma", {"italic": True}),
            " prohibition.  The express terms of the contract require that any "
            "change to scope or price be in a writing signed by the "
            "Superintendent or other officer the Board of Trustees has "
            "authorized.  (Compl., ¶ 11 & Ex. A, art. 12.2.)  Plaintiff’s "
            "implied-covenant theory asks the Court to impose an obligation on "
            "the District to pay for verbally-directed work that, by the "
            "contract’s express terms, is not compensable.  That is precisely "
            "the kind of extra-contractual obligation ",
            ("Guz", {"italic": True}),
            " forbids.  (",
            ("Id.", {"italic": False}),  # CORRUPT(LS-CITE-06c) — italics stripped from Id.
            " at pp. 349-350.)",
        ],
    )
    add_body_paragraph(
        doc,
        [
            "Plaintiff’s fallback allegation — that the District acted in "
            "bad faith by “refusing to process” the five proposed written "
            "change orders — fares no better.  Nothing in the contract "
            "obligated the District’s Superintendent or Board of Trustees to "
            "approve any particular change order, and the Complaint alleges no "
            "facts suggesting the District’s decision not to approve them was "
            "pretextual or untethered from a legitimate exercise of contract "
            "administration.  An implied-covenant claim premised on the "
            "non-exercise of a discretionary contractual right fails where, as "
            "here, the complaint pleads no facts that would take the defendant’s "
            "conduct outside the range of contractually-permitted behavior.  (",
            ("Carma", {"italic": True}),
            ", 2 Cal.4th at p. 374.)",
        ],
    )
    # CORRUPT(LS-SP-02) @ ¶65 — "sustained.  Because" -> "sustained. Because" (single space)
    add_body_paragraph(
        doc,
        [
            "The demurrer to the first cause of action should be sustained. "  # CORRUPT(LS-SP-02)
            "Because no amendment could cure the mismatch between what the "
            "written contract requires and what Plaintiff’s theory demands, "
            "leave to amend is not warranted.  (",
            ("Schifando", {"italic": True}),
            ", 31 Cal.4th at p. 1081.)",
        ],
    )

    # IV.B
    add_blank_line(doc)
    add_heading(
        doc,
        "B.  THE SECOND CAUSE OF ACTION FAILS BECAUSE VERBAL CHANGE ORDERS ARE NOT ENFORCEABLE AGAINST A SCHOOL DISTRICT.",
    )
    # CORRUPT(LS-CAP-05) @ ¶68 — "Public Contract Code section 20118.4 requires"
    # -> "Public Contract Code Section 20118.4 requires" (mid-sentence Section capitalized)
    add_body_paragraph(
        doc,
        [
            "The second cause of action pleads breach of contract premised "
            "entirely on alleged “verbal change orders” given by a District "
            "maintenance employee.  (Compl., ¶¶ 33-40.)  The claim fails "
            "as a matter of law for two independent reasons:  (1) Public Contract "
            "Code Section 20118.4 requires change orders on school-district "  # CORRUPT(LS-CAP-05)
            "public works to be in writing; and (2) Education Code sections "
            "17604 and 17605 require board approval — or validly-delegated "
            "approval — for contract amendments, and no such approval is, or "
            "could be, alleged."
        ],
    )
    # CORRUPT(LS-CITE-02) @ ¶69 — body-text "Public Contract Code section 20118.4 governs"
    # -> "Public Contract Code § 20118.4 governs" (outside-parens § symbol)
    add_body_paragraph(
        doc,
        [
            "Public Contract Code § 20118.4 governs change orders on "  # CORRUPT(LS-CITE-02)
            "public-works projects undertaken by school districts.  It provides "
            "a statutory framework under which a governing board may, by written "
            "change order, authorize the contractor to perform additional or "
            "changed work within defined limits.  The statute contemplates — "
            "and California courts have consistently enforced — a written "
            "change-order regime; it does not authorize verbal modifications.  "
            "(See Pub. Cont. Code, § 20118.4.)  The Court of Appeal has "
            "repeatedly confirmed that written change-order requirements in "
            "public-works contracts are to be strictly enforced, and that a "
            "contractor who proceeds without compliance does so at its own risk.  "
            "(",
            ("P&D Consultants, Inc. v. City of Carlsbad", {"italic": True}),
            " (2010) 190 Cal.App.4th 1332, 1340-1343.)",
        ],
    )
    # CORRUPT(LS-CITE-06g) @ ¶70 — change trailing "(Id.)" -> "(Ibid.)"
    # LS forbids ibid.; this is the canonical "Ibid." violation.
    add_body_paragraph(
        doc,
        [
            "Education Code sections 17604 and 17605 complete the statutory "
            "picture.  Section 17604 provides that contracts made by a school "
            "district are not valid or enforceable against the district unless "
            "approved or ratified by the governing board.  (Ed. Code, § "
            "17604.)  Section 17605 permits the governing board "
            "to delegate contracting authority to a district officer, but only "
            "by formal written resolution, and only within the limits the board "
            "specifies.  (Ed. Code, § 17605.)  A person who has not received "
            "a valid delegation under section 17605 has no authority to bind the "
            "district.  (",
            ("Ibid.", {"italic": True}),  # CORRUPT(LS-CITE-06g) — was Id.
            ")",
        ],
    )
    add_body_paragraph(
        doc,
        [
            "Plaintiff’s Complaint concedes the fatal point:  Mr. Ortiz, the "
            "Director of Maintenance and Operations, is alleged to have been "
            "the “on-site District representative most familiar with the "
            "Project” — nothing more.  (Compl., ¶ 15.)  The Complaint "
            "does not allege that the Board of Trustees delegated change-order "
            "authority to Mr. Ortiz under section 17605, or that the Board "
            "approved or ratified any of Mr. Ortiz’s alleged verbal "
            "directions.  Absent such allegations, no enforceable contract "
            "modification arose from Mr. Ortiz’s statements — however well-"
            "intentioned or practical his instructions may have been.  (Ed. "
            "Code, §§ 17604, 17605.)",
        ],
    )
    # CORRUPT(LS-CITE-HAL #17a) @ ¶72 — transpose digits in a real case's page number.
    # "City of Long Beach v. Mansell (1970) 3 Cal.3d 462, 493" -> "(1970) 3 Cal.3d 462, 439"
    # (the correct volume is 462; the 493 pin-cite was swapped to 439 — case exists, pin won't resolve)
    add_body_paragraph(
        doc,
        [
            "California decisions going back more than a century have applied "
            "these principles unforgivingly.  Strict compliance with "
            "statutory contracting requirements is treated as a jurisdictional "
            "prerequisite to the formation of a valid contract with a public "
            "entity, and contracts entered without the requisite statutory "
            "authority are void ",
            ("ab initio", {"italic": True}),
            ".  (",
            ("G.L. Mezzetta, Inc. v. City of American Canyon", {"italic": True}),
            " (2000) 78 Cal.App.4th 1087, 1092-1093; ",
            ("First Street Plaza Partners v. City of Los Angeles", {"italic": True}),
            " (1998) 65 Cal.App.4th 650, 668-669.)  And it is settled that a "
            "public agency cannot be estopped to assert the invalidity of a "
            "contract executed in excess of its statutory authority.  (",
            ("Air Quality Products, Inc. v. State of California", {"italic": True}),
            " (1979) 96 Cal.App.3d 340, 349; ",
            ("Lentz v. McMahon", {"italic": True}),
            " (1989) 49 Cal.3d 393, 399-400; see also ",
            ("City of Long Beach v. Mansell", {"italic": True}),
            " (1970) 3 Cal.3d 462, 439 [estoppel against government available "  # CORRUPT(LS-CITE-HAL) — pin 493 transposed to 439
            "only in unusual circumstances and never where it would defeat a "
            "strong public policy].)",
        ],
    )
    add_body_paragraph(
        doc,
        [
            "The Supreme Court has also foreclosed the very theory Plaintiff "
            "advances here — that a contractor may recover, on a total-cost "
            "or abandonment theory, for extras performed outside the written "
            "change-order regime.  (",
            ("Amelco Electric v. City of Thousand Oaks", {"italic": True}),
            " (2002) 27 Cal.4th 228, 234-235.)  The rule could not be clearer:  "
            "when a public contract contains express change-order provisions, "
            "those provisions must be honored, and a contractor’s course-of-"
            "performance attempt to bypass them is not a legally cognizable "
            "theory of recovery.  (",
            ("Id.", {"italic": True}),
            " at p. 235; see also ",
            ("Janis v. California State Lottery Com.", {"italic": True}),
            " (1998) 68 Cal.App.4th 824, 829-831 [estoppel against a public "
            "entity is rare and narrowly limited].)",
        ],
    )
    # CORRUPT(LS-SP-07) @ ¶74 — drop the Oxford comma in the three-item list
    # "no such officer, no such authorization, and no such consent" ->
    # "no such officer, no such authorization and no such consent"
    add_body_paragraph(
        doc,
        [
            "Every essential element of the contract claim fails at the pleading "
            "stage.  A contract requires mutual consent.  (Civ. Code, §§ "
            "1550, 1565.)  Here, the “consent” on the District side could "
            "only have come from an officer the Board of Trustees had authorized "
            "to give it — and Plaintiff’s Complaint alleges no such "
            "officer, no such authorization and no such consent.  The second "  # CORRUPT(LS-SP-07) — Oxford comma dropped
            "cause of action should be dismissed without leave to amend."
        ],
    )

    # IV.C
    add_blank_line(doc)
    add_heading(
        doc,
        "C.  THE THIRD CAUSE OF ACTION FAILS BECAUSE UNJUST ENRICHMENT IS NOT AVAILABLE AGAINST A PUBLIC ENTITY FOR WORK OUTSIDE AN AUTHORIZED CONTRACT.",
    )
    # CORRUPT(LS-CITE-HAL #17b) @ ¶77 — insert a wholly fake case citation.
    # "Lakeside School Dist. v. Coastal Builders, Inc. (2014) 229 Cal.App.4th 517"
    # is a fabricated name not present on fixtures/seed-citations.verified.md.
    add_body_paragraph(
        doc,
        [
            "Plaintiff’s third cause of action seeks the reasonable value of "
            "the additional work under a theory of unjust enrichment or quantum "
            "meruit.  (Compl., ¶¶ 44-51.)  That theory has been foreclosed "
            "against public entities in California for more than a century.  "
            "(See ",
            ("Lakeside School Dist. v. Coastal Builders, Inc.", {"italic": True}),  # CORRUPT(LS-CITE-HAL) — fabricated case
            " (2014) 229 Cal.App.4th 517, 523-524.)"
        ],
    )
    # CORRUPT(LS-ETAL-01) @ ¶78 — introduce "Reams v. Cooley, et al." (comma before et al.)
    # LS requires "et al." with no preceding comma (e.g., "Reams et al. v. Cooley") — the
    # prohibited form under LS-ETAL-01 is the comma-preceded variant.
    add_body_paragraph(
        doc,
        [
            "The rule originated in ",
            ("Reams v. Cooley, et al.", {"italic": True}),  # CORRUPT(LS-ETAL-01) — "; et al." with preceding comma
            " (1915) 171 Cal. 150, 154-156, in which the California Supreme "
            "Court held that a contractor who performed work under a contract "
            "that had not been validly authorized by the school district could "
            "not recover on a quantum-meruit theory for the reasonable value of "
            "the work performed.  The ",
            ("Reams", {"italic": True}),
            " rule was reaffirmed in ",
            ("Miller v. McKinnon", {"italic": True}),
            " (1942) 20 Cal.2d 83, 88-89, which held that a contract made in "
            "violation of a statutory mandate governing public-entity "
            "contracting is void and that no recovery in quantum meruit is "
            "available for services rendered under such a contract.",
        ],
    )
    # CORRUPT(LS-CITE-08d) @ ¶79 — add "supra" to a case short cite.
    # "First Street Plaza Partners, 65 Cal.App.4th at pp. 668-669" ->
    # "First Street Plaza Partners, supra, 65 Cal.App.4th at pp. 668-669"
    add_body_paragraph(
        doc,
        [
            "Modern decisions have applied the ",
            ("Reams/Miller", {"italic": True}),
            " rule across the full range of public-entity contracting contexts.  "
            "(",
            ("Katsura v. City of San Buenaventura", {"italic": True}),
            " (2007) 155 Cal.App.4th 104, 109-110 [no quantum-meruit recovery "
            "against city for services under a void contract]; ",
            ("First Street Plaza Partners", {"italic": True}),
            ", ",
            ("supra", {"italic": True}),  # CORRUPT(LS-CITE-08d) — "supra" injected into short cite
            ", 65 Cal.App.4th at pp. 668-669 [contracts made without statutory "
            "authority are void ",
            ("ab initio", {"italic": True}),
            " and confer no right to quasi-contractual recovery]; ",
            ("G.L. Mezzetta", {"italic": True}),
            ", 78 Cal.App.4th at pp. 1092-1093 [same].)  The rationale is "
            "straightforward:  the statutory contracting requirements exist to "
            "protect the public fisc, and permitting quantum-meruit recovery "
            "around those requirements would nullify an important statutory "
            "policy adopted for the benefit of the public.  (",
            ("Lentz", {"italic": True}),
            ", 49 Cal.3d at pp. 399-400.)",
        ],
    )
    # CORRUPT(LS-QUOTE-01) @ ¶80 — move the closing period from inside to outside
    # the quotation mark on one of the quoted labels.
    # "quantum meruit."  ->  "quantum meruit".
    add_body_paragraph(
        doc,
        [
            "Plaintiff cannot escape the rule by recharacterizing its claim as "
            "“unjust enrichment” rather than “quantum meruit”.  "  # CORRUPT(LS-QUOTE-01) — period outside the close-quote
            "California courts have applied the ",
            ("Reams/Miller", {"italic": True}),
            " bar without regard to the quasi-contractual label the plaintiff "
            "chooses.  (",
            ("Katsura", {"italic": True}),
            ", 155 Cal.App.4th at p. 109.)  Nor can Plaintiff rely on Government "
            "Code section 53060, which authorizes certain special-services "
            "contracts:  that section does not dispense with the board-approval "
            "and written-change-order requirements of the Education and Public "
            "Contract Codes.  (Gov. Code, § 53060.)  And public-entity tort "
            "liability, which also requires a statutory basis, provides no "
            "refuge either.  (Gov. Code, § 815, subd. (a).)",
        ],
    )
    # CORRUPT(LS-CITE-07) @ ¶81 — add "; et seq." with preceding comma.
    # "(See Pub. Cont. Code, § 22300.)" -> "(See Pub. Cont. Code, § 22300, et seq.)"
    add_body_paragraph(
        doc,
        [
            "Finally, statutory retention-substitution and kindred "
            "public-works protections — which exist, in part, to ensure that "
            "contractors are paid promptly for properly-authorized work — "
            "do not create a backdoor route to payment for unauthorized work.  "
            "(See Pub. Cont. Code, § 22300, et seq.)  Those protections "  # CORRUPT(LS-CITE-07) — comma before "et seq."
            "operate within the universe of authorized public contracts; they "
            "do not expand that universe."
        ],
    )
    add_body_paragraph(
        doc,
        [
            "Because the unjust-enrichment claim is foreclosed as a matter of "
            "law on the face of the Complaint, and because no amendment can "
            "transform verbally-directed work into a validly-authorized public "
            "contract, the demurrer to the third cause of action should be "
            "sustained without leave to amend.  (",
            ("Schifando", {"italic": True}),
            ", 31 Cal.4th at p. 1081.)",
        ],
    )

    # V. CONCLUSION
    add_blank_line(doc)
    add_heading(doc, "V.  CONCLUSION")
    add_body_paragraph(
        doc,
        [
            "For the foregoing reasons, the District respectfully requests that "
            "the Court sustain the District’s demurrer to the first, second, "
            "and third causes of action of the Complaint, without leave to "
            "amend, and grant such further relief as the Court deems just and "
            "proper."
        ],
    )

    # Signature block
    add_blank_line(doc)
    add_body_paragraph(
        doc,
        [
            "Dated:  April 22, 2026",
            "".ljust(20),
            ("LOZANO SMITH", {"bold": True}),
        ],
        first_line_indent=False,
    )
    add_blank_line(doc)
    add_body_paragraph(
        doc,
        [
            "".ljust(44),
            "By:  ",
            ("Margaret A. Delacroix", {"bold": True}),
        ],
        first_line_indent=False,
    )
    add_body_paragraph(
        doc,
        [
            "".ljust(44),
            "Margaret A. Delacroix",
        ],
        first_line_indent=False,
    )
    add_body_paragraph(
        doc,
        [
            "".ljust(44),
            "Jonathan R. Pemberton",
        ],
        first_line_indent=False,
    )
    add_body_paragraph(
        doc,
        [
            "".ljust(44),
            "Attorneys for Defendant",
        ],
        first_line_indent=False,
    )
    add_body_paragraph(
        doc,
        [
            "".ljust(44),
            "EASTVALE JOINT UNION HIGH",
        ],
        first_line_indent=False,
    )
    add_body_paragraph(
        doc,
        [
            "".ljust(44),
            "SCHOOL DISTRICT",
        ],
        first_line_indent=False,
    )


# --------------------------------------------------------------------------- #
# Default-style setup
# --------------------------------------------------------------------------- #


def configure_default_styles(doc) -> None:
    """Set the Normal style defaults so any stray paragraph uses TNR 12pt."""
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = FONT_SIZE
    # east-asian / cs fallback on the style element
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rFonts.set(qn(f"w:{attr}"), FONT_NAME)

    # Page margins: 1 inch all around (typical CA superior court).
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)


# --------------------------------------------------------------------------- #
# Main
# --------------------------------------------------------------------------- #


def main() -> None:
    doc = Document()
    configure_default_styles(doc)

    build_caption(doc)
    build_notice(doc)
    build_memorandum(doc)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"Wrote {OUTPUT_PATH} ({OUTPUT_PATH.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
