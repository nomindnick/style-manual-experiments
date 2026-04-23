"""Minimal `.docx` loader.

Wraps python-docx so rules don't depend on its internals directly. Exposes only
what a Phase-0/1 rule needs: full text, paragraphs with text + style metadata,
and enough hooks to let Tier 2 rules inspect font / spacing / indent later.

The scope here is deliberately minimal — add accessors when the first rule that
needs them lands. See PLAN.md "Open questions" #1.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterator

from docx import Document as _DocxDocument
from docx.document import Document as _DocxDocumentType


@dataclass
class Paragraph:
    """One paragraph of a document.

    Style metadata is represented as plain values where python-docx makes that
    possible; None indicates the value is inherited from a parent style (which
    is normal for most docs and not a problem).
    """

    index: int
    text: str
    style_name: str | None
    font_name: str | None
    font_size_pt: float | None
    line_spacing: float | None  # may be ratio (1.5) or absolute pt (24.0) depending on rule
    line_spacing_rule: str | None  # 'EXACTLY', 'MULTIPLE', etc.
    first_line_indent_in: float | None
    left_indent_in: float | None
    right_indent_in: float | None


class Document:
    """A loaded `.docx` document.

    Intentionally does not inherit from python-docx's Document type; we want the
    freedom to change representation later without breaking rule code.
    """

    def __init__(self, docx_doc: _DocxDocumentType, source_path: Path | None = None):
        self._doc = docx_doc
        self.source_path = source_path
        self._paragraphs: list[Paragraph] | None = None

    @classmethod
    def load(cls, path: str | Path) -> "Document":
        p = Path(path)
        return cls(_DocxDocument(str(p)), source_path=p)

    @property
    def paragraphs(self) -> list[Paragraph]:
        if self._paragraphs is None:
            self._paragraphs = list(self._build_paragraphs())
        return self._paragraphs

    @property
    def text(self) -> str:
        """Full document text with paragraphs joined by newlines."""
        return "\n".join(p.text for p in self.paragraphs)

    def _build_paragraphs(self) -> Iterator[Paragraph]:
        for i, para in enumerate(self._doc.paragraphs):
            pf = para.paragraph_format
            # python-docx returns Pt-like objects; .pt gives a float
            font_size_pt = (
                para.style.font.size.pt
                if para.style and para.style.font and para.style.font.size
                else None
            )
            # Prefer run-level font size if the paragraph's first run has one
            if para.runs and para.runs[0].font.size is not None:
                font_size_pt = para.runs[0].font.size.pt

            line_spacing = pf.line_spacing
            # python-docx returns either a float multiplier (e.g. 1.5) or an Emu/Pt
            # depending on the rule. Normalize if it's a length.
            line_spacing_value: float | None
            if line_spacing is None:
                line_spacing_value = None
            elif isinstance(line_spacing, (int, float)):
                line_spacing_value = float(line_spacing)
            else:  # Length object with .pt
                line_spacing_value = float(getattr(line_spacing, "pt", line_spacing))

            line_spacing_rule = (
                pf.line_spacing_rule.name if pf.line_spacing_rule is not None else None
            )

            def _inches(length):
                return float(length.inches) if length is not None else None

            yield Paragraph(
                index=i,
                text=para.text,
                style_name=para.style.name if para.style else None,
                font_name=(para.runs[0].font.name if para.runs and para.runs[0].font.name else None),
                font_size_pt=font_size_pt,
                line_spacing=line_spacing_value,
                line_spacing_rule=line_spacing_rule,
                first_line_indent_in=_inches(pf.first_line_indent),
                left_indent_in=_inches(pf.left_indent),
                right_indent_in=_inches(pf.right_indent),
            )

    @classmethod
    def from_text(cls, text: str) -> "Document":
        """Build a synthetic Document from plain text, one paragraph per line.

        Useful for rule unit tests that don't need real .docx styling. Style
        metadata is all None — rules that need style info should test against
        real fixtures, not synthetic text.
        """
        doc = _DocxDocument()
        for line in text.split("\n"):
            doc.add_paragraph(line)
        return cls(doc)
